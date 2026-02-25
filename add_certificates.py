"""
证书和关系证明添加方案

需求：在投标文件中添加公司关系证明、资质证书等PDF内容
方式：图片形式或直接插入PDF

## 现有情况

1. **数据结构**：
   - qualifications.json 中有 cert_file 字段，记录了证书文件路径
   - data/certificates/ 目录下有很多证书PDF文件

2. **当前实现**（generator.py 第608-720行）：
   - 会显示证书名称和等级
   - 会显示证书文件路径（文字形式）
   - 但不会插入PDF内容或图片

## 简单解决方案（推荐）

### 方案1：直接插入PDF文件（最简单）

**优点**：
- Word 支持直接插入PDF文件
- 不需要转换
- 代码简单
- 用户可以双击打开PDF查看完整内容

**缺点**：
- 不是以图片形式显示
- 需要双击才能打开

**实现**：
```python
from docx.shared import Inches

def add_certificates_simple(doc: Document, qualifications: List[Dict], data_dir: Path):
    """
    添加证书到文档（简单版 - 直接插入PDF）
    
    Args:
        doc: Word 文档对象
        qualifications: 资质列表
        data_dir: 数据目录
    """
    p = doc.add_paragraph()
    run = p.add_run("第3章 企业资质")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    
    doc.add_paragraph()
    
    # 显示每个证书
    for qual in qualifications:
        if not qual.get('cert_file'):
            continue
        
        cert_path = data_dir / qual['cert_file']
        if not cert_path.exists():
            continue
        
        # 添加证书名称
        p = doc.add_paragraph()
        run = p.add_run(f"• {qual['name']}（{qual['level']}）")
        run.bold = True
        
        # 添加证书文件路径
        doc.add_paragraph(f"  证书文件：{qual['cert_file']}")
        
        # 【新增】直接插入PDF文件
        try:
            doc.add_paragraph()  # 空行
            doc.add_picture(str(cert_path), width=Inches(6.0))  # 插入PDF，宽度6英寸
            doc.add_paragraph()  # 空行
        except Exception as e:
            # 如果插入失败，显示文件路径
            doc.add_paragraph(f"  （证书文件详见附件：{qual['cert_file']}）")
```

### 方案2：显示文件路径 + 提供下载（最稳妥）

**优点**：
- 代码最简单
- 不会出现格式问题
- 可以添加多个文件
- 用户可以选择性查看

**缺点**：
- 不是直接显示在文档中
- 需要额外操作才能查看

**实现**：
```python
def add_certificates_with_link(doc: Document, qualifications: List[Dict], data_dir: Path, output_dir: Path):
    """
    添加证书到文档（带下载链接）
    
    Args:
        doc: Word 文档对象
        qualifications: 资质列表
        data_dir: 数据目录
        output_dir: 输出目录
    """
    p = doc.add_paragraph()
    run = p.add_run("第3章 企业资质")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    
    doc.add_paragraph()
    
    # 创建证书文件列表
    cert_files_to_copy = []
    
    for qual in qualifications:
        if not qual.get('cert_file'):
            continue
        
        cert_path = data_dir / qual['cert_file']
        if not cert_path.exists():
            continue
        
        cert_files_to_copy.append(cert_path)
        
        # 添加证书信息
        p = doc.add_paragraph()
        run = p.add_run(f"• {qual['name']}（{qual['level']}）")
        run.bold = True
        
        doc.add_paragraph(f"  证书文件：{qual['cert_file']}")
    
    # 添加说明
    doc.add_paragraph()
    doc.add_paragraph("提示：以上证书文件已复制到文档附件文件夹，可双击打开查看。")
    
    return cert_files_to_copy
```

### 方案3：PDF 转图片（最完整，但较复杂）

**优点**：
- 以图片形式显示，效果最好
- 不需要双击就能看到
- 可以控制显示大小

**缺点**：
- 需要安装 pdf2image 库
- 需要安装 Ghostscript
- 转换可能较慢
- 文件大小会增加

**实现**：
```python
def add_certificates_as_image(doc: Document, qualifications: List[Dict], data_dir: Path):
    """
    添加证书到文档（PDF 转图片）
    
    Args:
        doc: Word 文档对象
        qualifications: 资质列表
        data_dir: 数据目录
    """
    from pdf2image import convert_from_path
    import tempfile
    
    p = doc.add_paragraph()
    run = p.add_run("第3章 企业资质")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    
    doc.add_paragraph()
    
    for qual in qualifications:
        if not qual.get('cert_file'):
            continue
        
        cert_path = data_dir / qual['cert_file']
        if not cert_path.exists():
            continue
        
        # 添加证书名称
        p = doc.add_paragraph()
        run = p.add_run(f"• {qual['name']}（{qual['level']}）")
        run.bold = True
        
        # 转换PDF为图片
        try:
            # 只转换第一页
            with tempfile.TemporaryDirectory() as temp_dir:
                images = convert_from_path(
                    str(cert_path),
                    output_folder=temp_dir,
                    first_page_only=True,
                    dpi=150  # 分辨率
                )
                
                if images:
                    # 插入图片
                    doc.add_picture(str(images[0]), width=Inches(6.0))
                    doc.add_paragraph()
        except Exception as e:
            # 如果转换失败，显示文件路径
            doc.add_paragraph(f"  （证书文件详见附件：{qual['cert_file']}）")
            doc.add_paragraph(f"  （转换失败：{str(e)}）")
```

## 推荐方案

**建议使用方案1**：
- 直接插入PDF文件
- 不需要额外安装库
- Word 支持良好
- 代码简单

**如果后续需要更好的显示效果**，可以升级到方案3（PDF转图片）。

## 使用方法

```python
# 修改 generator.py 中的 _add_qualifications_v2 方法

# 修改前：
for cert in system_certs[:10]:  # 最多显示10个
    doc.add_paragraph(f"• {cert['name']}（{cert['level']}）")
    if cert.get('cert_file'):
        doc.add_paragraph(f"  证书文件：{cert['cert_file']}")

# 修改后（方案1）：
for cert in system_certs[:10]:
    if cert.get('cert_file'):
        cert_path = self.data_dir / cert['cert_file']
        if cert_path.exists():
            # 插入PDF文件
            doc.add_picture(str(cert_path), width=Inches(6.0))
            doc.add_paragraph()
```

## 注意事项

1. **文件路径问题**：
   - 证书文件路径需要正确（相对于 data_dir）
   - 建议使用绝对路径或确保 data_dir 正确

2. **显示大小**：
   - width=Inches(6.0) 表示宽度为6英寸
   - 可以根据需要调整

3. **文件数量**：
   - 证书文件太多可能导致文档过大
   - 建议限制显示数量（如只显示前10个）

4. **兼容性**：
   - 某些旧版本 Word 可能不支持插入PDF
   - 建议测试后确认效果
"""

if __name__ == "__main__":
    print(__doc__)
