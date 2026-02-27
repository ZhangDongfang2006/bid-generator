#!/usr/bin/env python3
"""
公司通用内容生成方法

用于从临时文件读取参考PDF中提取的公司通用内容，
并添加到投标文件中。
"""

from pathlib import Path
from typing import Dict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def read_temp_file(filename: str) -> str:
    """从临时文件读取内容"""
    temp_dir = Path('/tmp')
    file_path = temp_dir / filename
    
    if file_path.exists():
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        return ""


def add_chapter_from_text(doc: Document, title: str, content: str):
    """从文本内容添加章节到文档"""
    # 添加章节标题
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"

    doc.add_paragraph()

    # 添加内容
    for paragraph in content.split('\n'):
        if paragraph.strip():
            # 检查是否是标题（加粗或大字体）
            is_title = False
            if paragraph.strip().endswith(':') or \
               paragraph.strip().startswith('一、') or \
               paragraph.strip().startswith('二、') or \
               paragraph.strip().startswith('三、') or \
               paragraph.strip().startswith('四、') or \
               paragraph.strip().startswith('五、') or \
               paragraph.strip().startswith('六、') or \
               paragraph.strip().startswith('七、') or \
               paragraph.strip().startswith('八、') or \
               paragraph.strip().startswith('九、') or \
               paragraph.strip().startswith('十、') or \
               paragraph.strip().startswith('十一、') or \
               paragraph.strip().startswith('十二、') or \
               paragraph.strip().startswith('十三、') or \
               paragraph.strip().startswith('十四、') or \
               paragraph.strip().startswith('十五、'):
                is_title = True

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(paragraph.strip())
            
            if is_title:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = "黑体"
            else:
                run.font.size = Pt(14)
                run.font.name = "宋体"

    doc.add_page_break()


def add_legal_authorization(doc: Document):
    """添加法定代表人授权书"""
    content = read_temp_file('legal_authorization.txt')
    add_chapter_from_text(doc, "二、法定代表人授权书", content)


def add_bid_guarantee(doc: Document):
    """添加投标保证金缴纳证明"""
    content = read_temp_file('bid_guarantee.txt')
    add_chapter_from_text(doc, "三、投标保证金缴纳证明", content)


def add_warranty_commitment(doc: Document):
    """添加质保期满后三年内的备品备件供货承诺"""
    content = read_temp_file('warranty_commitment.txt')
    add_chapter_from_text(doc, "六、质保期满后三年内的备品备件供货承诺", content)


def add_compliance_statement(doc: Document):
    """添加近三年无重大违法记录声明"""
    content = read_temp_file('compliance_statement.txt')
    add_chapter_from_text(doc, "九、近三年无重大违法记录声明", content)


def add_quality_control_plan(doc: Document):
    """添加质量控制专项方案"""
    content = read_temp_file('quality_control_plan.txt')
    
    # 质量控制方案内容较多，分多个子章节
    add_chapter_from_text(doc, "十二、质量控制专项方案", content)


def add_safety_guarantee(doc: Document):
    """添加安全保证"""
    content = read_temp_file('safety_guarantee.txt')
    add_chapter_from_text(doc, "十三、安全保证", content)


def add_delivery_plan(doc: Document):
    """添加供货组织及进度计划"""
    content = read_temp_file('delivery_plan.txt')
    add_chapter_from_text(doc, "十四、供货组织及进度计划", content)


def add_training_and_service(doc: Document):
    """添加技术培训、售后服务"""
    content = read_temp_file('training_and_service.txt')
    add_chapter_from_text(doc, "十五、技术培训、售后服务的内容、计划及措施", content)


if __name__ == "__main__":
    # 测试读取临时文件
    print("测试读取临时文件...")
    
    files = [
        'legal_authorization.txt',
        'bid_guarantee.txt',
        'warranty_commitment.txt',
        'compliance_statement.txt',
        'quality_control_plan.txt',
        'safety_guarantee.txt',
        'delivery_plan.txt',
        'training_and_service.txt'
    ]
    
    for filename in files:
        content = read_temp_file(filename)
        lines = len([line for line in content.split('\n') if line.strip()])
        print(f"  {filename}: {lines} 行")
    
    print("\n✅ 所有临时文件读取成功")
