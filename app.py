"""
投标文件生成应用 - AI置信度分析版
"""

import streamlit as st
import os
from pathlib import Path
from io import BytesIO
from datetime import datetime
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入本地模块
from parser import TenderParser, ParseResult
from generator import BidDocumentGenerator as BidGenerator
from database import CompanyDatabase
import config


# ==================== 配置 ====================

# 页面设置
st.set_page_config(
    page_title="海越（湖北）电气 - 智能投标文件生成系统",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS
st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
}
.stMarkdown {
    color: white;
}
</style>
""", unsafe_allow_html=True)

# ==================== 初始化 ====================

# 初始化数据库
data_dir = Path(__file__).parent / "data"
db = CompanyDatabase(data_dir)

# 初始化解析器
parser = TenderParser(data_dir)

# 初始化生成器
templates_dir = Path(__file__).parent / "templates"
generator = BidGenerator(db, templates_dir)

# ==================== 会话状态 ====================

# 初始化session state
if 'tender_info' not in st.session_state:
    st.session_state.tender_info = {}

if 'matched_data' not in st.session_state:
    st.session_state.matched_data = {}

if 'parse_result' not in st.session_state:
    st.session_state.parse_result = None

if 'bid_generated' not in st.session_state:
    st.session_state.bid_generated = False

if 'preview_doc_bytes' not in st.session_state:
    st.session_state.preview_doc_bytes = None

if 'preview_available' not in st.session_state:
    st.session_state.preview_available = False

# ==================== 主界面 ====================

# 侧边栏
with st.sidebar:
    st.header("🏢 海越投标助手")
    st.divider()
    
    # 资料管理入口
    if st.button("📊 资料管理", use_container_width=True):
        st.session_state.active_page = "data_management"
    
    # 快速操作
    st.divider()
    st.markdown("### 🚀 快速操作")
    
    # 查看数据库状态
    st.markdown("#### 📊 数据库状态")
    st.markdown(f"- **资质**: {len(db.get_qualifications())} 项")
    st.markdown(f"- **案例**: {len(db.get_cases())} 项")
    st.markdown(f"- **产品**: {len(db.get_products())} 项")
    st.markdown(f"- **人员**: {len(db.get_personnel())} 项")

# 主内容区
if st.session_state.get('active_page') == 'data_management':
    st.title("📊 资料管理")
    st.markdown("请在本地文件系统中管理以下目录中的内容：")
    st.markdown(f"- `{data_dir}/qualifications.json`")
    st.markdown(f"- `{data_dir}/cases.json`")
    st.markdown(f"- `{data_dir}/products.json`")
    st.markdown(f"- `{data_dir}/personnel.json`")
    
else:
    st.title("📄 智能投标文件生成")
    st.markdown("---")
    
    # 第一步：上传招标文件
    st.header("📤 第一步：上传招标文件")
    st.markdown("支持 PDF、Word (.docx, .doc) 格式的招标文件")
    
    # 文件上传
    uploaded_file = st.file_uploader(
        "上传招标文件",
        type=['pdf', 'docx', 'doc'],
        help="支持 PDF、Word 格式",
        key="tender_file_uploader"
    )
    
    # 解析上传的文件
    if uploaded_file is not None:
        # 保存到临时文件
        temp_file = Path("temp") / uploaded_file.name
        temp_file.parent.mkdir(exist_ok=True)
        
        with open(temp_file, 'wb') as f:
            f.write(uploaded_file.getbuffer())
        
        # 解析文件
        st.info("🔄 正在解析文件...")
        parse_result = parser.parse_file(temp_file)
        st.session_state.parse_result = parse_result
        
        # 显示解析结果
        st.markdown("---")
        st.subheader("📋 文件解析结果")
        
        # 显示置信度
        st.markdown(f"### {parse_result.get_confidence_color()} 解析置信度")
        st.metric(
            "置信度",
            f"{parse_result.confidence_score:.2f}",
            delta=f"{parse_result.confidence_score:.2f}",
            help=f"{parse_result.confidence_level} - AI 对文件解析的可信程度"
        )
        
        # 显示解析出的需求
        st.markdown(f"**提取需求**: {len(parse_result.requirements)}")
        
        # 显示需求列表
        with st.expander("📝 查看提取的需求", expanded=False):
            for i, req in enumerate(parse_result.requirements, 1):
                st.text(f"{i}. {req}")
        
        # 提供人工校验
        st.markdown("---")
        st.subheader("🔍 人工校验")
        st.markdown("如果解析结果有误，可以在下方修改：")
        
        # 编辑需求
        edited_requirements = []
        for i, req in enumerate(parse_result.requirements, 1):
            edited_req = st.text_area(
                f"需求 {i}",
                value=req,
                key=f"req_edit_{i}",
                height=50
            )
            edited_requirements.append(edited_req if edited_req else req)
        
        # 更新session state
        st.session_state.tender_info['requirements'] = edited_requirements
        
        # 显示建议
        if parse_result.confidence_level != "高":
            st.warning(f"⚠️ {parse_result.confidence_level}置信度：建议仔细校验解析结果")
            
            # 生成改进建议
            suggestions = parser._get_suggestions(parse_result)
            if suggestions:
                st.markdown("---")
                st.subheader("💡 改进建议")
                for i, suggestion in enumerate(suggestions, 1):
                    st.markdown(f"{i}. {suggestion}")
    
    # 第二步：匹配公司数据
    if st.session_state.parse_result:
        st.markdown("---")
        st.header("🎯 第二步：匹配公司数据")
        st.markdown("根据提取的需求，智能匹配公司的资质、案例、产品")
        
        # 匹配资质
        st.subheader("📋 匹配资质")
        requirements = st.session_state.tender_info.get('requirements', [])
        matched_qualifications = db.match_qualifications(requirements)
        
        st.markdown(f"**匹配结果**: {len(matched_qualifications)} 项")
        
        # 显示匹配的资质（最多前5个）
        with st.expander("查看匹配的资质", expanded=False):
            for i, qual in enumerate(matched_qualifications[:5], 1):
                st.markdown(f"{i}. **{qual['name']}** - {qual['level']}")
                st.caption(f"证书编号：{qual.get('cert_no', 'N/A')}")
        
        # 更新session state
        st.session_state.matched_data['qualifications'] = matched_qualifications
        
        # 匹配案例
        st.subheader("📋 匹配案例")
        matched_cases = db.match_cases(requirements)
        
        st.markdown(f"**匹配结果**: {len(matched_cases)} 项")
        
        # 显示匹配的案例（最多前5个）
        with st.expander("查看匹配的案例", expanded=False):
            for i, case in enumerate(matched_cases[:5], 1):
                st.markdown(f"{i}. **{case['project_name']}**")
                st.caption(f"客户：{case.get('client', 'N/A')} | 金额：{case.get('amount', 0):,.0f} 元")
        
        # 更新session state
        st.session_state.matched_data['cases'] = matched_cases
        
        # 匹配产品
        st.subheader("📋 匹配产品")
        matched_products = db.match_products(requirements)
        
        st.markdown(f"**匹配结果**: {len(matched_products)} 项")
        
        # 显示匹配的产品（最多前5个）
        with st.expander("查看匹配的产品", expanded=False):
            for i, product in enumerate(matched_products[:5], 1):
                st.markdown(f"{i}. **{product['name']}**")
                st.caption(f"型号：{product['model']} | 分类：{product.get('category', 'N/A')}")
        
        # 更新session state
        st.session_state.matched_data['products'] = matched_products
    
    # 第三步：生成投标文件
    if st.session_state.matched_data:
        st.markdown("---")
        st.header("🚀 第三步：生成投标文件")
        st.markdown("一键生成技术标和商务标（或合并文档）")
        
        # 显示证书图片选项
        show_cert_images = st.checkbox("显示证书图片", value=True, key="show_cert_images")
        st.caption("勾选后，生成的投标文件中将包含证书图片（PDF转图片）")
        
        # 生成选项
        col1, col2 = st.columns(2)
        
        with col1:
            separate_bids = st.checkbox("技术标和商务标分开生成", value=True, key="separate_bids")
            st.caption("勾选后，将生成两个独立的文件")
        
        with col2:
            preview_first = st.checkbox("生成前预览（推荐）", value=True, key="preview_first")
            st.caption("勾选后，先生成预览版本，再确认下载正式版本")
        
        # 生成按钮
        if st.button("🚀 生成投标文件", type="primary", key="generate_bid"):
            try:
                st.info("🔄 正在生成投标文件...")
                
                # 更新 tender_info
                st.session_state.tender_info['show_cert_images'] = show_cert_images
                st.session_state.tender_info['generate_time'] = datetime.now().isoformat()
                
                # 准备匹配数据
                matched_data = st.session_state.matched_data
                
                # 调试信息
                st.write(f"生成信息：")
                st.write(f"  - 显示证书图片：{show_cert_images}")
                st.write(f"  - 匹配资质：{len(matched_data.get('qualifications', []))}")
                st.write(f"  - 匹配案例：{len(matched_data.get('cases', []))}")
                st.write(f"  - 匹配产品：{len(matched_data.get('products', []))}")
                
                if separate_bids:
                    # 生成技术标和商务标分开
                    if preview_first:
                        # 生成预览版本（简化内容）
                        output_paths = generator.generate_separate_bids_preview(
                            st.session_state.tender_info,
                            config.COMPANY_INFO,
                            matched_data
                        )
                        st.success("✅ 预览文件生成成功！")
                        st.session_state.preview_available = True
                    else:
                        # 生成完整版本
                        output_paths = generator.generate_separate_bids(
                            st.session_state.tender_info,
                            config.COMPANY_INFO,
                            matched_data
                        )
                        st.success("✅ 投标文件生成成功！")
                        st.session_state.bid_generated = True
                else:
                    # 生成单一文件
                    if preview_first:
                        # 生成预览版本
                        output_path = generator.generate_bid_preview(
                            st.session_state.tender_info,
                            config.COMPANY_INFO,
                            matched_data
                        )
                        st.success("✅ 预览文件生成成功！")
                        st.session_state.preview_available = True
                    else:
                        # 生成完整版本
                        output_path = generator.generate_bid(
                            st.session_state.tender_info,
                            config.COMPANY_INFO,
                            matched_data
                        )
                        st.success("✅ 投标文件生成成功！")
                        st.session_state.bid_generated = True
                
                # 显示下载链接
                if separate_bids:
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        if output_paths and output_paths[0]:
                            tech_bid_path = output_paths[0]
                            with open(tech_bid_path, 'rb') as f:
                                st.download_button(
                                    label="📥 下载技术标",
                                    data=f,
                                    file_name=f"技术标_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                    
                    with col2:
                        if output_paths and output_paths[1]:
                            biz_bid_path = output_paths[1]
                            with open(biz_bid_path, 'rb') as f:
                                st.download_button(
                                    label="📥 下载商务标",
                                    data=f,
                                    file_name=f"商务标_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                else:
                    if output_path:
                        with open(output_path, 'rb') as f:
                            st.download_button(
                                label="📥 下载投标文件",
                                data=f,
                                file_name=f"投标文件_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                
            except Exception as e:
                st.error(f"❌ 生成失败：{e}")
                st.markdown(f"**错误详情**：{str(e)}")
    
    # 第四步：预览投标文件
    if st.session_state.preview_available:
        st.markdown("---")
        st.header("👁 第四步：预览投标文件")
        st.markdown("在浏览器中预览生成的投标文件，无需下载")
        
        # 读取生成的文件
        output_dir = Path("output")
        if output_dir.exists():
            files = list(output_dir.glob("*.docx"))
            
            if files:
                latest_file = max(files, key=lambda f: f.stat().st_mtime)
                
                try:
                    # 读取文档
                    from docx import Document
                    doc = Document(str(latest_file))
                    
                    # 在浏览器中预览
                    st.markdown("### 📄 预览内容")
                    
                    # 显示文档标题
                    for para in doc.paragraphs[:5]:
                        if para.text.strip():
                            st.markdown(f"**{para.text}**")
                    
                    # 显示文档内容预览
                    with st.expander("查看更多内容", expanded=False):
                        for para in doc.paragraphs[5:20]:
                            if para.text.strip():
                                st.text(para.text)
                    
                    # 下载预览版本
                    with open(latest_file, 'rb') as f:
                        st.download_button(
                            label="📥 下载预览版本",
                            data=f,
                            file_name=f"预览_{latest_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # 生成正式版本按钮
                    if st.button("🚀 确认并生成正式版本", type="primary", key="generate_final"):
                        try:
                            st.info("🔄 正在生成正式版本...")
                            
                            # 使用完整的数据生成
                            matched_data = st.session_state.matched_data
                            st.session_state.tender_info['generate_time'] = datetime.now().isoformat()
                            st.session_state.tender_info['is_final'] = True
                            
                            # 生成完整版本
                            if st.session_state.tender_info.get('separate_bids'):
                                output_paths = generator.generate_separate_bids(
                                    st.session_state.tender_info,
                                    config.COMPANY_INFO,
                                    matched_data
                                )
                            else:
                                output_path = generator.generate_bid(
                                    st.session_state.tender_info,
                                    config.COMPANY_INFO,
                                    matched_data
                                )
                            
                            st.success("✅ 正式版本生成成功！")
                            st.session_state.preview_available = False
                            st.session_state.bid_generated = True
                            
                            # 提示下载
                            st.info("📥 请在下方下载正式版本")
                            
                        except Exception as e:
                            st.error(f"❌ 生成失败：{e}")
                    
                except Exception as e:
                    st.warning(f"⚠️ 预览失败：{e}")
    
    # 第五步：下载投标文件
    if st.session_state.bid_generated:
        st.markdown("---")
        st.header("📥 第五步：下载投标文件")
        
        # 显示生成状态
        st.markdown(f"**生成时间**: {st.session_state.tender_info.get('generate_time', 'N/A')}")
        st.markdown(f"**是否正式版本**: {'是' if st.session_state.tender_info.get('is_final') else '预览版本'}")
        
        # 查找生成的文件
        output_dir = Path("output")
        if output_dir.exists():
            files = list(output_dir.glob("*.docx"))
            
            if files:
                latest_file = max(files, key=lambda f: f.stat().st_mtime)
                file_size = latest_file.stat().st_size / 1024  # KB
                
                # 显示文件信息
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**文件名**")
                    st.code(latest_file.name, language="text")
                
                with col2:
                    st.markdown("**文件大小**")
                    st.metric("", f"{file_size:.1f} KB")
                
                # 下载按钮
                with open(latest_file, 'rb') as f:
                    st.download_button(
                        label="📥 下载最新版本",
                        data=f,
                        file_name=latest_file.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("⚠️ 没有找到生成的文件")
    
    # 底部信息
    st.markdown("---")
    st.markdown("---")
    st.markdown("### ℹ️ 使用说明")
    st.markdown("""
    1. 上传招标文件（PDF/Word）
    2. 查看AI解析置信度
    3. 校验提取的需求
    4. 查看智能匹配结果
    5. 生成投标文件
    6. 在浏览器中预览（推荐）
    7. 下载正式版本
    
    **注意事项**：
    - 生成的投标文件为Word格式
    - 可以在 Microsoft Word 或 WPS 中打开
    - 建议在提交前人工检查所有内容
    - 如有问题，请联系技术支持
    """)
