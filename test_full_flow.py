#!/usr/bin/env python3
"""
测试完整的PDF转图片并插入Word的流程
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches
import json
from pdf2image import convert_from_path
from PIL import Image as PILImage
import tempfile
import uuid

# 读取资质数据
data_dir = Path('/Users/zhangdongfang/workspace/bid-generator/data')
with open(data_dir / 'qualifications.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 只测试前3个证书
test_certs = data.get('qualifications', [])[:3]

print(f"测试 {len(test_certs)} 个证书...")
print()

# 创建临时目录用于保存图片
temp_dir = Path(tempfile.mkdtemp(prefix='test_bid_'))
print(f"临时目录: {temp_dir}")
print()

# 创建Word文档
doc = Document()
p = doc.add_paragraph()
run = p.add_run("测试 - 证书图片插入")
run.bold = True
run.font.size = 20

converted_images = {}

# 转换每个证书PDF为图片
for i, cert in enumerate(test_certs, 1):
    if not cert.get('cert_file'):
        print(f"跳过证书 {i}: 没有PDF文件")
        continue

    cert_path = data_dir / cert['cert_file']
    print(f"处理证书 {i}: {cert['name']}")
    print(f"  PDF路径: {cert_path}")

    if not cert_path.exists():
        print(f"  ✗ 文件不存在")
        continue

    # 转换PDF为图片
    try:
        print(f"  开始转换PDF...")
        images = convert_from_path(
            str(cert_path),
            first_page=1,
            last_page=1,
            dpi=200,
            fmt='jpg',
            use_cropbox=True
        )

        if images:
            # 调整图片大小
            img = images[0]
            img_width, img_height = img.size
            print(f"  原始大小: {img_width} x {img_height}")

            target_width = 500
            if img_width > target_width:
                ratio = target_width / img_width
                new_height = int(img_height * ratio)
                img = img.resize((target_width, new_height), PILImage.LANCZOS)
                print(f"  调整后大小: {target_width} x {new_height}")

            # 保存到临时目录
            img_filename = f"{cert['id']}_{uuid.uuid4().hex[:8]}.jpg"
            img_path = temp_dir / img_filename
            img.save(img_path, quality=85)
            print(f"  保存图片: {img_path}")
            print(f"  图片大小: {img_path.stat().st_size} bytes")

            converted_images[cert['id']] = {
                'name': cert['name'],
                'level': cert['level'],
                'images': [img_path]
            }
            print(f"  ✓ 转换成功")
        else:
            print(f"  ✗ 转换失败: 没有返回图片")

    except Exception as e:
        print(f"  ✗ 转换失败: {e}")
        import traceback
        traceback.print_exc()

    print()

print(f"总共转换成功: {len(converted_images)} 个证书")
print()

# 插入图片到Word文档
print("插入图片到Word文档...")

for cert_id, cert_data in converted_images.items():
    doc.add_paragraph()
    doc.add_paragraph(f"证书: {cert_data['name']}（{cert_data['level']}）")

    for img_path in cert_data['images']:
        try:
            print(f"  插入图片: {img_path}")
            doc.add_paragraph()
            doc.add_picture(str(img_path), width=Inches(5.5))
            doc.add_paragraph()
            print(f"  ✓ 插入成功")
        except Exception as e:
            print(f"  ✗ 插入失败: {e}")
            doc.add_paragraph(f"  （图片插入失败: {e}）")

print()

# 保存Word文档
output_path = Path('/tmp/test_certificate_images.docx')
doc.save(output_path)
print(f"✓ Word文档已保存: {output_path}")
print()

# 清理临时目录
import shutil
shutil.rmtree(temp_dir, ignore_errors=True)
print(f"✓ 已清理临时目录")

print()
print("=" * 60)
print("测试完成！")
print(f"请打开 Word 文档查看图片: {output_path}")
