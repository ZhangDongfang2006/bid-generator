#!/usr/bin/env python3
"""
测试完整的资质证书转换流程
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
import json
from pdf2image import convert_from_path
from PIL import Image as PILImage
import tempfile
import uuid

# 读取资质数据
data_dir = Path('/Users/zhangdongfang/workspace/bid-generator/data')
with open(data_dir / 'qualifications.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

qualifications = data.get('qualifications', [])

print(f"总共有 {len(qualifications)} 个证书")
print()

# 按分类统计
system_certs = [q for q in qualifications if "体系" in q.get("name", "") or "认证" in q.get("name", "")]
credit_certs = [q for q in qualifications if "AAA" in q.get("name", "") or "信用" in q.get("name", "")]
honor_certs = [q for q in qualifications if "重点" in q.get("name", "") or "质量奖" in q.get("name", "")]
partner_certs = [q for q in qualifications if "授权" in q.get("name", "") or "合作" in q.get("name", "")]

print(f"3.1 体系认证证书: {len(system_certs)} 个")
print(f"3.2 信用等级证书: {len(credit_certs)} 个")
print(f"3.3 重点荣誉证书: {len(honor_certs)} 个")
print(f"3.4 合作伙伴证书: {len(partner_certs)} 个")
print()

# 统计有 PDF 的证书
with_cert_file = [q for q in qualifications if q.get('cert_file')]
print(f"有 PDF 文件的证书: {len(with_cert_file)} 个")
print()

# 转换所有有 PDF 的证书
temp_dir = Path(tempfile.mkdtemp(prefix='test_full_certs_'))
print(f"临时目录: {temp_dir}")
print()

converted_images = {}
total = 0
success = 0
failed = 0

for i, cert in enumerate(qualifications, 1):
    if not cert.get('cert_file'):
        continue
    
    cert_path = data_dir / cert['cert_file']
    if not cert_path.exists():
        print(f"✗ 证书文件不存在: {cert['cert_file']}")
        failed += 1
        total += 1
        continue
    
    # 转换PDF为图片
    try:
        images = convert_from_path(
            str(cert_path),
            first_page=1,
            last_page=1,
            dpi=200,
            fmt='jpg',
            use_cropbox=True
        )
        
        if images:
            img = images[0]
            img_width, img_height = img.size
            
            target_width = 500
            if img_width > target_width:
                ratio = target_width / img_width
                new_height = int(img_height * ratio)
                img = img.resize((target_width, new_height), PILImage.LANCZOS)
            
            img_filename = f"{cert['id']}_{uuid.uuid4().hex[:8]}.jpg"
            img_path = temp_dir / img_filename
            img.save(img_path, quality=85)
            
            converted_images[cert['id']] = {
                'name': cert['name'],
                'level': cert['level'],
                'images': [img_path]
            }
            success += 1
            print(f"✓ {i}. {cert['name']} 转换成功")
        else:
            converted_images[cert['id']] = {
                'name': cert['name'],
                'level': cert['level'],
                'images': [],
                'error': '转换失败'
            }
            failed += 1
            print(f"✗ {i}. {cert['name']} 转换失败")
    except Exception as e:
        print(f"✗ {i}. {cert['name']} 转换失败: {e}")
        converted_images[cert['id']] = {
            'name': cert['name'],
            'level': cert['level'],
            'images': [],
            'error': str(e)
        }
        failed += 1
    
    total += 1

print()
print(f"转换完成: 总计 {total}, 成功 {success}, 失败 {failed}")
print()

# 创建Word文档
doc = Document()
p = doc.add_paragraph()
run = p.add_run("测试 - 所有证书图片")
run.bold = True
run.font.size = 20

doc.add_paragraph()

# 3.1 体系认证证书
print("插入 3.1 体系认证证书...")
p = doc.add_paragraph()
run = p.add_run("3.1 体系认证证书")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

for cert in system_certs:
    images = converted_images.get(cert['id'], {}).get('images', [])
    
    p = doc.add_paragraph()
    run = p.add_run(f"• {cert['name']}（{cert['level']}）")
    run.bold = True
    
    for img_path in images:
        try:
            doc.add_paragraph()
            doc.add_picture(str(img_path), width=Inches(5.5))
            doc.add_paragraph()
            print(f"  ✓ {cert['name']} 图片插入成功")
        except Exception as e:
            doc.add_paragraph(f"  （图片插入失败: {e}）")
            print(f"  ✗ {cert['name']} 图片插入失败: {e}")

# 3.5 其他证书
print()
print("插入 3.5 其他证书...")
classified_ids = set(cert['id'] for cert in system_certs + credit_certs + honor_certs + partner_certs)
other_certs_with_images = [q for q in qualifications if q['id'] not in classified_ids and q.get('cert_file')]

if other_certs_with_images:
    p = doc.add_paragraph()
    run = p.add_run("3.5 其他证书")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    for cert in other_certs_with_images:
        images = converted_images.get(cert['id'], {}).get('images', [])

        p = doc.add_paragraph()
        run = p.add_run(f"• {cert['name']}")
        run.bold = True

        for img_path in images:
            try:
                doc.add_paragraph()
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.add_paragraph()
                print(f"  ✓ {cert['name']} 图片插入成功")
            except Exception as e:
                doc.add_paragraph(f"  （图片插入失败: {e}）")
                print(f"  ✗ {cert['name']} 图片插入失败: {e}")

# 保存Word文档
output_path = Path('/tmp/test_all_certs_images.docx')
doc.save(output_path)

# 清理临时目录
import shutil
shutil.rmtree(temp_dir, ignore_errors=True)

print()
print("=" * 60)
print(f"✓ Word文档已保存: {output_path}")
print(f"✓ 已清理临时目录")
print()
print(f"统计:")
print(f"  - 总证书: {len(qualifications)} 个")
print(f"  - 有PDF: {len(with_cert_file)} 个")
print(f"  - 转换成功: {success} 个")
print(f"  - 转换失败: {failed} 个")
print()
print("请打开 Word 文档查看所有证书图片")
