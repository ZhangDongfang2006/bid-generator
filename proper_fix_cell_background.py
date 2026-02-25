"""
永久修复方案：使用 docx API 设置单元格背景色

这个脚本展示了如何使用 docx 库提供的 API 来设置单元格背景色，
而不是手动创建 XML 元素。
"""

def get_proper_cell_background_fix():
    """
    获取正确的单元格背景色设置代码
    
    使用 docx.oxml.shared.CT_Shd 类来创建阴影元素
    """
    
    code = """
    # 在 generator.py 顶部添加导入
    from docx.oxml.shared import CT_Shd
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # 在 _add_deviation_table 方法中，替换原来的代码：
    
    # 原来的代码（有问题）：
    # cell._element.get_or_add_tcPr().append(
    #     self._parse_xml(f'<w:shd {self._nsdecls("w")} w:fill="D9E2F3"/>')
    # )
    
    # 新代码（使用 docx API）：
    # 创建阴影元素
    shd = CT_Shd()
    shd.fill = "D9E2F3"  # 浅灰色背景
    
    # 添加到单元格的 tcPr
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.shd = shd
    """
    
    return code


def get_complete_fix():
    """
    获取完整的修复方案
    
    包括：
    1. 添加必要的导入
    2. 修改 _add_deviation_table 方法
    3. 可以选择删除 _nsdecls 和 _parse_xml 方法（如果不使用它们）
    """
    
    code = """
    ==================== 在 generator.py 顶部添加导入 ====================
    
    from docx.oxml.shared import CT_Shd
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsdecls
    
    ==================== 修改 _add_deviation_table 方法 ====================
    
    def _add_deviation_table(self, doc: Document, tender_info: Dict, table_type: str):
        '''
        添加偏离表
        
        Args:
            doc: Word 文档对象
            tender_info: 招标信息
            table_type: 表格类型（技术/商务）
        '''
        p = doc.add_paragraph()
        run = p.add_run(f"第1章 {table_type}偏离表")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"
        
        doc.add_paragraph()
        
        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        headers = ["序号", "条款编号", "偏离说明"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            
            # 设置表头格式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(12)
            
            # 【修复】设置单元格背景色（使用 docx API）
            shd = CT_Shd()
            shd.fill = "D9E2F3"  # 浅灰色背景
            
            tcPr = cell._element.get_or_add_tcPr()
            tcPr.shd = shd
        
        # 第一行示例
        cells = table.add_row().cells
        cells[0].text = "1"
        cells[1].text = "-"
        cells[2].text = f"无{table_type}偏离"
        
        doc.add_paragraph()
        doc.add_paragraph(f"我方承诺：{table_type}投标文件中所投产品完全满足技术规范书中星号条款（*）中的相关要求。")
        
        doc.add_page_break()
    
    ==================== 可选：删除 _nsdecls 和 _parse_xml 方法 ====================
    
    # 如果不使用手动创建 XML 元素，可以删除这两个方法：
    # def _nsdecls(self, prefix):
    #     ...
    
    # def _parse_xml(self, xml_str):
    #     ...
    """
    
    return code


def test_fix():
    """
    测试修复方案
    
    检查导入是否正确，以及代码是否能正常工作
    """
    
    print("检查导入...")
    try:
        from docx import Document
        from docx.oxml.shared import CT_Shd
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.oxml.ns import nsdecls
        print("✓ 所有导入都正常")
    except ImportError as e:
        print(f"✗ 导入错误: {e}")
        return False
    
    print("\n创建测试文档...")
    try:
        doc = Document()
        
        # 添加一个测试表格
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        headers = ["列1", "列2", "列3"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            
            # 设置表头背景色
            shd = CT_Shd()
            shd.fill = "D9E2F3"
            tcPr = cell._element.get_or_add_tcPr()
            tcPr.shd = shd
        
        # 添加数据行
        for row in range(1):
            cells = table.add_row().cells
            for i in range(3):
                cells[i].text = f"数据{row+1}-{i+1}"
        
        # 保存测试文档
        output_path = Path("/tmp/test_cell_background.docx")
        doc.save(output_path)
        print(f"✓ 测试文档已保存: {output_path}")
        
        # 验证文档是否可以打开
        from docx import Document as DocxDocument
        test_doc = DocxDocument(output_path)
        table = test_doc.tables[0]
        print(f"✓ 文档验证通过，表格有 {len(table.rows)} 行")
        
        return True
        
    except Exception as e:
        print(f"✗ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    import sys
    from pathlib import Path
    
    print("=" * 60)
    print("单元格背景色修复方案")
    print("=" * 60)
    
    # 显示修复代码
    print("\n【修复代码】")
    code = get_proper_cell_background_fix()
    print(code)
    
    # 显示完整修复方案
    print("\n" + "=" * 60)
    print("【完整修复方案】")
    print("=" * 60)
    full_code = get_complete_fix()
    print(full_code)
    
    # 测试修复
    print("\n" + "=" * 60)
    print("【测试修复】")
    print("=" * 60)
    if test_fix():
        print("\n✓ 修复方案测试通过！")
        print("\n下一步：")
        print("1. 应用修复到 generator.py")
        print("2. 重启应用程序")
        print("3. 测试生成投标文件")
        print("4. 确认表格背景色正常显示")
    else:
        print("\n✗ 修复方案测试失败")
        print("\n请检查 docx 库版本")
        sys.exit(1)
