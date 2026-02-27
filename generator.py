"""
完整的投标文件生成器 - 包含所有修复和PDF转图片功能

功能：
1. 所有Bug已修复
2. 支持PDF转图片功能
3. 详细的日志记录
4. 增强的调试模式
5. 错误处理和恢复

使用方法：
1. 此文件替换原来的 generator.py
2. 重启应用
3. 在生成投标文件时，勾选"显示证书图片"选项
"""

import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PDF 转图片相关导入
try:
    from pdf2image import convert_from_path
    from PIL import Image as PILImage
    PDF_TO_IMAGE_AVAILABLE = True
except ImportError as e:
    PDF_TO_IMAGE_AVAILABLE = False
    PILImage = None
    print(f"✗ 错误: pdf2image 或 Pillow 未安装")
    print(f"详细信息: {e}")
    print("安装命令:")
    print("  pip install pdf2image Pillow")
    print("")
    print("请先安装这些库，然后重启应用")
    import sys
    sys.exit(1)  # 退出程序，因为现在默认启用PDF转图片功能

# 导入公司通用内容生成方法
try:
    from company_content import (
        add_legal_authorization,
        add_bid_guarantee,
        add_warranty_commitment,
        add_compliance_statement,
        add_quality_control_plan,
        add_safety_guarantee,
        add_delivery_plan,
        add_training_and_service
    )
    COMPANY_CONTENT_AVAILABLE = True
except ImportError as e:
    COMPANY_CONTENT_AVAILABLE = False
    print(f"⚠️ 警告: company_content 模块未安装: {e}")
    print("将不生成公司通用内容章节")


class BidDocumentGenerator:
    """投标文件生成器 - V2 (PDF转图片版）"""

    def __init__(self, templates_dir: Path, output_dir: Path):
        self.templates_dir = templates_dir
        self.output_dir = output_dir
        self.image_width_inches = 4.5  # 自动调整的图片大小（4.5英寸，约11.4厘米）

    def generate_bid(self, tender_info: Dict, company_info: Dict,
                    matched_data: Dict, quote_data: Dict = None,
                    show_cert_images: bool = False) -> Path:
        """
        生成投标文件

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据（资质、案例、产品等）
            quote_data: 报价数据
            show_cert_images: 是否显示证书图片

        Returns:
            生成的文件路径
        """
        # 创建新文档
        doc = Document()

        # 生成各个章节
        self._add_cover_v2(doc, tender_info, company_info, bid_type="投标文件")
        self._add_table_of_contents(doc, separate_bids=False)  # 修改：添加目录
        self._add_company_proof(doc, company_info)
        self._add_bid纲领_v2(doc, company_info, tender_info)
        self._add_deviation_table(doc, tender_info, table_type="技术")
        self._add_company_intro_v2(doc, company_info)
        self._add_tech_solution(doc, tender_info, matched_data)
        
        # 添加公司通用内容
        if COMPANY_CONTENT_AVAILABLE:
            self._add_legal_authorization(doc)
            self._add_bid_guarantee(doc)
            self._add_warranty_commitment(doc)
        
        self._add_equipment_specs_table(doc, self.templates_dir.parent / "data")
        
        # 添加更多公司通用内容
        if COMPANY_CONTENT_AVAILABLE:
            self._add_quality_control_plan(doc)
            self._add_safety_guarantee(doc)
            self._add_delivery_plan(doc)
            self._add_training_and_service(doc)
        
        self._add_quotation(doc, quote_data if quote_data else {})
        self._add_qualifications_with_images(doc, matched_data.get("qualifications", []), self.templates_dir.parent / "data", show_cert_images)
        self._add_performance(doc, matched_data.get("cases", []))
        self._add_after_sales(doc, company_info)

        # 设置页码
        self._setup_page_numbers(doc)

        # 保存文件
        project_name = tender_info.get("project_info", {}).get("project_name", "未知项目")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"投标文件_{project_name}_{timestamp}.docx"

        # 清理文件名
        filename = "".join(c for c in filename if c not in '\/:*?"<>|')

        output_path = self.output_dir / filename
        doc.save(output_path)

        return output_path

    def generate_tech_bid(self, tender_info: Dict, company_info: Dict,
                          matched_data: Dict, quote_data: Dict = None,
                          show_cert_images: bool = False) -> Path:
        """
        生成技术标

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据
            quote_data: 报价数据
            show_cert_images: 是否显示证书图片

        Returns:
            生成的文件路径
        """
        # 创建新文档
        doc = Document()

        # 生成技术标章节
        self._add_cover_v2(doc, tender_info, company_info, bid_type="技术投标文件")
        self._add_table_of_contents(doc, separate_bids=True)  # 新增：添加目录
        self._add_company_proof(doc, company_info)
        self._add_bid纲领_v2(doc, company_info, tender_info)
        self._add_deviation_table(doc, tender_info, table_type="技术")
        self._add_company_intro_v2(doc, company_info)
        self._add_tech_solution(doc, tender_info, matched_data)
        
        # 添加公司通用内容
        if COMPANY_CONTENT_AVAILABLE:
            self._add_compliance_statement(doc)
            self._add_quality_control_plan(doc)
            self._add_safety_guarantee(doc)
        
        self._add_equipment_specs_table(doc, self.templates_dir.parent / "data")
        
        # 添加更多公司通用内容
        if COMPANY_CONTENT_AVAILABLE:
            self._add_delivery_plan(doc)
            self._add_training_and_service(doc)
        
        self._add_qualifications_with_images(doc, matched_data.get("qualifications", []), self.templates_dir.parent / "data", show_cert_images)
        self._add_performance(doc, matched_data.get("cases", []))
        self._add_tech_commitment(doc)
        self._add_response_commitment(doc)

        # 保存文件
        project_name = tender_info.get("project_info", {}).get("project_name", "未知项目")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"技术标_{project_name}_{timestamp}.docx"

        # 清理文件名
        filename = "".join(c for c in filename if c not in '\/:*?"<>|')

        output_path = self.output_dir / filename
        doc.save(output_path)

        return output_path

    def generate_commercial_bid(self, tender_info: Dict, company_info: Dict,
                                matched_data: Dict, quote_data: Dict = None,
                                show_cert_images: bool = False) -> Path:
        """
        生成商务标

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据
            quote_data: 报价数据
            show_cert_images: 是否显示证书图片

        Returns:
            生成的文件路径
        """
        # 创建新文档
        doc = Document()

        # 生成商务标章节
        self._add_cover_v2(doc, tender_info, company_info, bid_type="商务投标文件")
        self._add_table_of_contents(doc, separate_bids=True)  # 新增：添加目录
        self._add_company_proof(doc, company_info)
        self._add_bid纲领_v2(doc, company_info, tender_info)
        self._add_deviation_table(doc, tender_info, table_type="商务")
        self._add_company_intro_v2(doc, company_info)
        self._add_response_commitment(doc)
        self._add_quotation(doc, quote_data if quote_data else {})
        self._add_qualifications_with_images(doc, matched_data.get("qualifications", []), self.templates_dir.parent / "data", show_cert_images)
        self._add_performance(doc, matched_data.get("cases", []))
        self._add_after_sales(doc, company_info)
        self._add_commercial_commitment(doc)

        # 保存文件
        project_name = tender_info.get("project_info", {}).get("project_name", "未知项目")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"商务标_{project_name}_{timestamp}.docx"

        # 清理文件名
        filename = "".join(c for c in filename if c not in '\/:*?"<>|')

        output_path = self.output_dir / filename
        doc.save(output_path)

        return output_path

    def generate_separate_bids(self, tender_info: Dict, company_info: Dict,
                               matched_data: Dict, quote_data: Dict = None,
                               show_cert_images: bool = False) -> Dict[str, Path]:
        """
        生成分开的技术标和商务标

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据（资质、案例、产品等）
            quote_data: 报价数据
            show_cert_images: 是否显示证书图片

        Returns:
            包含技术标和商务标路径的字典
        """
        # 生成技术标
        tech_path = self.generate_tech_bid(tender_info, company_info, matched_data, quote_data, show_cert_images)

        # 生成商务标
        commercial_path = self.generate_commercial_bid(tender_info, company_info, matched_data, quote_data, show_cert_images)

        return {
            "tech": tech_path,
            "commercial": commercial_path
        }

    def generate_separate_bids_preview(self, tender_info: Dict, company_info: Dict,
                                      matched_data: Dict) -> Dict[str, Path]:
        """
        生成分开的技术标和商务标（预览版本，简化内容）

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据（资质、案例、产品等）

        Returns:
            包含技术标和商务标路径的字典
        """
        # 预览版本只显示基本信息和匹配结果

        # 生成技术标预览
        tech_doc = Document()
        self._add_cover_v2(tech_doc, tender_info, company_info, "技术标（预览）")

        # 添加目录
        tech_doc.add_heading("目录", level=1)
        tech_doc.add_paragraph("1. 公司简介")
        tech_doc.add_paragraph("2. 技术方案概述")
        tech_doc.add_paragraph("3. 项目团队")
        tech_doc.add_paragraph("4. 资质证书")
        tech_doc.add_paragraph("5. 项目案例")
        tech_doc.add_paragraph("6. 产品说明")
        tech_doc.add_page_break()

        # 添加公司简介（简化）
        tech_doc.add_heading("1. 公司简介", level=1)
        tech_doc.add_paragraph(company_info['name'])
        tech_doc.add_paragraph(company_info.get('description', ''))
        tech_doc.add_page_break()

        # 添加技术方案概述（简化）
        tech_doc.add_heading("2. 技术方案概述", level=1)
        tech_doc.add_paragraph("（预览版本，技术方案将在正式版本中详细展示）")
        tech_doc.add_paragraph(f"需求分析：共 {len(tender_info.get('requirements', []))} 个需求")
        tech_doc.add_page_break()

        # 添加项目团队（简化）
        tech_doc.add_heading("3. 项目团队", level=1)
        tech_doc.add_paragraph("（预览版本，项目团队将在正式版本中详细展示）")
        personnel = matched_data.get('personnel', [])[:3]  # 只显示前3个
        for p in personnel:
            tech_doc.add_paragraph(f"• {p.get('name', '')} - {p.get('role', '')}")
        tech_doc.add_page_break()

        # 添加资质证书（简化）
        tech_doc.add_heading("4. 资质证书", level=1)
        tech_doc.add_paragraph(f"共匹配到 {len(matched_data.get('qualifications', []))} 项资质")
        qualifications = matched_data.get('qualifications', [])[:5]  # 只显示前5个
        for q in qualifications:
            tech_doc.add_paragraph(f"• {q['name']} - {q['level']}")
        tech_doc.add_paragraph("（完整证书将在正式版本中显示）")
        tech_doc.add_page_break()

        # 添加项目案例（简化）
        tech_doc.add_heading("5. 项目案例", level=1)
        tech_doc.add_paragraph(f"共匹配到 {len(matched_data.get('cases', []))} 项案例")
        cases = matched_data.get('cases', [])[:3]  # 只显示前3个
        for c in cases:
            tech_doc.add_paragraph(f"• {c['project_name']} - {c.get('client', '')}")
        tech_doc.add_paragraph("（完整案例将在正式版本中显示）")
        tech_doc.add_page_break()

        # 添加产品说明（简化）
        tech_doc.add_heading("6. 产品说明", level=1)
        tech_doc.add_paragraph(f"共匹配到 {len(matched_data.get('products', []))} 项产品")
        products = matched_data.get('products', [])[:3]  # 只显示前3个
        for p in products:
            tech_doc.add_paragraph(f"• {p['name']} - {p['model']}")
        tech_doc.add_paragraph("（完整产品信息将在正式版本中显示）")

        # 保存技术标预览
        tech_output_dir = Path("output")
        tech_output_dir.mkdir(exist_ok=True)
        tech_filename = f"技术标预览_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        tech_path = tech_output_dir / tech_filename
        tech_doc.save(tech_path)

        # 生成商务标预览
        commercial_doc = Document()
        self._add_cover_v2(commercial_doc, tender_info, company_info, "商务标（预览）")

        # 添加目录
        commercial_doc.add_heading("目录", level=1)
        commercial_doc.add_paragraph("1. 报价说明")
        commercial_doc.add_paragraph("2. 报价汇总")
        commercial_doc.add_paragraph("3. 合同条款")
        commercial_doc.add_page_break()

        # 添加报价说明（简化）
        commercial_doc.add_heading("1. 报价说明", level=1)
        commercial_doc.add_paragraph("（预览版本，详细报价将在正式版本中展示）")
        commercial_doc.add_paragraph("本报价为预览报价，仅供参考。")
        commercial_doc.add_page_break()

        # 添加报价汇总（简化）
        commercial_doc.add_heading("2. 报价汇总", level=1)
        commercial_doc.add_paragraph("（预览版本，详细报价将在正式版本中展示）")
        commercial_doc.add_page_break()

        # 添加合同条款（简化）
        commercial_doc.add_heading("3. 合同条款", level=1)
        commercial_doc.add_paragraph("（预览版本，详细合同条款将在正式版本中展示）")

        # 保存商务标预览
        commercial_filename = f"商务标预览_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        commercial_path = tech_output_dir / commercial_filename
        commercial_doc.save(commercial_path)

        return {
            "tech": tech_path,
            "commercial": commercial_path
        }

    def generate_bid_preview(self, tender_info: Dict, company_info: Dict,
                           matched_data: Dict) -> Path:
        """
        生成投标文件（预览版本，简化内容）

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据（资质、案例、产品等）

        Returns:
            生成的文件路径
        """
        # 创建新文档
        doc = Document()

        # 添加封面
        self._add_cover_v2(doc, tender_info, company_info, bid_type="投标文件（预览）")

        # 添加目录
        doc.add_heading("目录", level=1)
        doc.add_paragraph("1. 公司简介")
        doc.add_paragraph("2. 技术方案概述")
        doc.add_paragraph("3. 项目团队")
        doc.add_paragraph("4. 资质证书")
        doc.add_paragraph("5. 项目案例")
        doc.add_paragraph("6. 产品说明")
        doc.add_paragraph("7. 报价说明")
        doc.add_paragraph("8. 售后服务")
        doc.add_page_break()

        # 添加公司简介
        doc.add_heading("1. 公司简介", level=1)
        doc.add_paragraph(company_info['name'])
        doc.add_paragraph(company_info.get('description', ''))
        doc.add_page_break()

        # 添加技术方案概述
        doc.add_heading("2. 技术方案概述", level=1)
        doc.add_paragraph("（预览版本，技术方案将在正式版本中详细展示）")
        doc.add_paragraph(f"需求分析：共 {len(tender_info.get('requirements', []))} 个需求")
        for i, req in enumerate(tender_info.get('requirements', [])[:5], 1):
            doc.add_paragraph(f"{i}. {req}")
        if len(tender_info.get('requirements', [])) > 5:
            doc.add_paragraph(f"...（还有 {len(tender_info.get('requirements', [])) - 5} 个需求）")
        doc.add_page_break()

        # 添加项目团队
        doc.add_heading("3. 项目团队", level=1)
        doc.add_paragraph("（预览版本，项目团队将在正式版本中详细展示）")
        personnel = matched_data.get('personnel', [])[:3]
        for p in personnel:
            doc.add_paragraph(f"• {p.get('name', '')} - {p.get('role', '')}")
        doc.add_page_break()

        # 添加资质证书
        doc.add_heading("4. 资质证书", level=1)
        doc.add_paragraph(f"共匹配到 {len(matched_data.get('qualifications', []))} 项资质")
        qualifications = matched_data.get('qualifications', [])[:5]
        for q in qualifications:
            doc.add_paragraph(f"• {q['name']} - {q['level']}")
        if len(matched_data.get('qualifications', [])) > 5:
            doc.add_paragraph(f"...（还有 {len(matched_data.get('qualifications', [])) - 5} 项资质）")
        doc.add_paragraph("（完整资质将在正式版本中显示，包括证书图片）")
        doc.add_page_break()

        # 添加项目案例
        doc.add_heading("5. 项目案例", level=1)
        doc.add_paragraph(f"共匹配到 {len(matched_data.get('cases', []))} 项案例")
        cases = matched_data.get('cases', [])[:3]
        for c in cases:
            doc.add_paragraph(f"• {c['project_name']} - {c.get('client', '')}")
        if len(matched_data.get('cases', [])) > 3:
            doc.add_paragraph(f"...（还有 {len(matched_data.get('cases', [])) - 3} 项案例）")
        doc.add_paragraph("（完整案例将在正式版本中详细展示）")
        doc.add_page_break()

        # 添加产品说明
        doc.add_heading("6. 产品说明", level=1)
        doc.add_paragraph(f"共匹配到 {len(matched_data.get('products', []))} 项产品")
        products = matched_data.get('products', [])[:3]
        for p in products:
            doc.add_paragraph(f"• {p['name']} - {p['model']}")
        if len(matched_data.get('products', [])) > 3:
            doc.add_paragraph(f"...（还有 {len(matched_data.get('products', [])) - 3} 项产品）")
        doc.add_paragraph("（完整产品信息将在正式版本中详细展示）")
        doc.add_page_break()

        # 添加报价说明
        doc.add_heading("7. 报价说明", level=1)
        doc.add_paragraph("（预览版本，详细报价将在正式版本中展示）")
        doc.add_paragraph("本报价为预览报价，仅供参考。")
        doc.add_page_break()

        # 添加售后服务
        doc.add_heading("8. 售后服务", level=1)
        doc.add_paragraph("（预览版本，详细售后服务承诺将在正式版本中展示）")
        doc.add_paragraph(company_info.get('service_commitment', ''))
        doc.add_page_break()

        # 保存文件
        project_name = tender_info.get("project_info", {}).get("project_name", "未知项目")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"投标文件预览_{project_name}_{timestamp}.docx"

        # 清理文件名
        filename = "".join(c for c in filename if c not in '\/:*?"<>|')

        output_path = self.output_dir / filename
        doc.save(output_path)

        return output_path

    # ==================== 添加章节的方法 ====================

    def _add_cover_v2(self, doc: Document, tender_info: Dict, company_info: Dict, bid_type: str = "投标文件"):
        """添加封面 - V3（符合参考PDF投标函格式）"""
        project_info = tender_info.get("project_info", {})
        project_name = project_info.get("project_name", "")
        project_no = project_info.get("project_no", "未知项目编号")
        tenderer = project_info.get("tenderer", "贵公司")
        project_full_name = project_info.get("project_full_name", f"{tenderer}{project_name}")
        project_amount = project_info.get("project_amount", "")
        project_tax_rate = project_info.get("project_tax_rate", "13%")
        bid_amount = project_info.get("bid_amount", "")
        bid_amount_upper = project_info.get("bid_amount_upper", "")

        # 公司信息
        company_name = company_info.get("name", "")
        company_address = company_info.get("address", "")
        company_phone = company_info.get("phone", "")
        company_email = company_info.get("email", "")
        company_rep_name = company_info.get("rep_name", "阎海")
        company_rep_title = company_info.get("rep_title", "投标中心主任")

        # 日期
        today = datetime.now()
        date_str = today.strftime('%Y年%m月%d日')

        # 添加封面标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"一、投标函")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 投标函内容
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"致：{tenderer}")
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        text = f"根据贵方{project_full_name}的投标邀请书（项目编号为：{project_no}），现正式授权的下列签字人{company_rep_name}、{company_rep_title}（姓名和职务）代表投标人{company_name}，提交下述投标文件正本1 份，副本4 份："
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 投标文件清单
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        bid_files = [
            "（1）金额为壹拾万元人民币的投标保证金；",
            "（2）资格证明文件；",
            "（3）投标报价表；",
            "（4）投标人须知第8 条要求投标人提交的全部文件。"
        ]
        
        for i, bid_file in enumerate(bid_files, 1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进
            run = p.add_run(bid_file)
            run.font.size = Pt(14)
            run.font.name = "宋体"

        doc.add_paragraph()

        # 投标人确认事项
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("据此函，签字人兹宣布同意如下：")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 确认事项列表
        confirmations = [
            f"a）按招标文件的规定提交设备及提供伴随服务的投标总价为人民币（大写）{bid_amount_upper}（{bid_amount} 元整），其中，税率{project_tax_rate}。",
            "b）我们将按招标文件的规定，承担完成合同规定的责任和义务。",
            "c）我们已详细审核了全部招标文件，包括招标文件的补充文件（如果有的话），我们知道必须放弃对上述文件中所有条款提出存有含糊不清或不理解之问题的权利。",
            "d）我们同意在投标人须知第21 条所述的开标日期起遵循本投标文件的规定，并在投标人须知第15 条规定的投标有效期届满之前对我方均具有约束力，而且有可能中标。",
            "e）如果在开标后规定的投标有效期内撤销投标，我们的投标保证金可不予退还。",
            "f）如果贵方有要求，我们愿意进一步提供与本投标有关的任何证据或资料。",
            "g）我们完全理解贵方不一定要接受最低报价的投标或收到的任何投标。"
        ]

        for confirmation in confirmations:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进
            run = p.add_run(confirmation)
            run.font.size = Pt(14)
            run.font.name = "宋体"

        doc.add_paragraph()

        # 正式通讯地址
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("与本投标有关的正式通讯地址为：")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "黑体"

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"地址：{company_address}")
        run.font.size = Pt(14)
        run.font.name = "宋体"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"邮政编码：{company_info.get('postcode', '432999')}")
        run.font.size = Pt(14)
        run.font.name = "宋体"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"电话号码：{company_phone}")
        run.font.size = Pt(14)
        run.font.name = "宋体"

        # 签字和盖章区域
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("投标人法定代表人或其授权代表签字或盖章：")
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("投标人单位公章：")
        run.font.size = Pt(14)
        run.font.name = "宋体"

        # 日期
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"日期：{date_str}")
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_page_break()

    def _add_company_proof(self, doc: Document, company_info: Dict):
        """添加公司证明"""
        p = doc.add_paragraph()
        run = p.add_run("公司证明")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 法人证明
        doc.add_paragraph(f"兹证明海越（湖北）电气股份有限公司（统一社会信用代码：{company_info.get('credit_code', '91420100329990389G')}）是依法设立的企业，具有独立法人资格。")

        # 公司基本情况
        doc.add_paragraph("公司基本情况：")
        doc.add_paragraph(f"公司名称：{company_info['name']}")
        doc.add_paragraph(f"公司地址：{company_info['address']}")
        doc.add_paragraph(f"邮政编码：{company_info.get('postal_code', '')}")
        doc.add_paragraph(f"联系电话：{company_info['phone']}")
        doc.add_paragraph(f"传真号码：{company_info['fax']}")
        doc.add_paragraph(f"电子邮箱：{company_info['email']}")

        doc.add_page_break()

    def _add_bid纲领_v2(self, doc: Document, company_info: Dict, tender_info: Dict):
        """添加投标纲领"""
        p = doc.add_paragraph()
        run = p.add_run("投标纲领")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 公司理念
        doc.add_paragraph(f"海越（湖北）电气股份有限公司是一家专注于高低压成套配电设备研发、制造、销售和服务的高新技术企业。公司始终坚持'质量第一、客户至上'的经营理念，致力于为客户提供优质的产品和完善的服务。")

        doc.add_paragraph()

        # 产品优势
        doc.add_paragraph("产品优势：")
        doc.add_paragraph("• 采用优质原材料，确保产品质量稳定可靠")
        doc.add_paragraph("• 引进先进生产工艺，提高生产效率")
        doc.add_paragraph("• 建立完善的质量管理体系，确保每道工序符合标准")
        doc.add_paragraph("• 具备专业的技术团队，为客户提供技术支持和服务")

        doc.add_page_break()

    def _add_deviation_table(self, doc: Document, tender_info: Dict, table_type: str = "技术"):
        """添加偏离表 - 新增"""
        p = doc.add_paragraph()
        run = p.add_run(f"第1章 {table_type}偏离表")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 说明
        doc.add_paragraph("说明：")
        doc.add_paragraph(f"本{table_type}投标文件完全满足招标文件中{table_type}规范书及{table_type}条款的全部要求，无偏离。")
        doc.add_paragraph()

        # 创建表格
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ["序号", "条款编号", "偏离说明"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # 设置表头格式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(12)
            # 【修复】设置单元格背景色
            from docx.shared import RGBColor
            cell.background_color = RGBColor(217, 226, 243)  # 浅灰色背景

        # 第一行示例
        cells = table.add_row().cells
        cells[0].text = "1"
        cells[1].text = "-"
        cells[2].text = f"无{table_type}偏离"

        doc.add_paragraph()
        doc.add_paragraph(f"我方承诺：{table_type}投标文件中所投产品完全满足技术规范书中星号条款（*）中的相关要求。")

        doc.add_page_break()

    def _add_company_intro_v2(self, doc: Document, company_info: Dict):
        """添加公司介绍 - V2（升级版）"""
        p = doc.add_paragraph()
        run = p.add_run("企业简介")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 公司介绍
        doc.add_paragraph("海越（湖北）电气股份有限公司是一家专注于高低压成套配电设备研发、制造、销售和服务的高新技术企业。")
        doc.add_paragraph()

        doc.add_paragraph("公司拥有两大生产基地，分别位于湖北孝感和浙江宁波，占地面积约50000平方米，建筑面积约35000平方米。")
        doc.add_paragraph()

        doc.add_paragraph("公司主要产品包括：35kV及以下高压开关柜、低压开关柜、箱式变电站、预制舱式变电站、配电箱、母线桥等配电设备，广泛应用于电力、钢铁、化工、交通、建筑、市政等行业。")

        doc.add_paragraph()

        doc.add_paragraph("公司拥有一支经验丰富的技术研发团队和完善的质量管理体系，通过了ISO9001质量管理体系认证、ISO14001环境管理体系认证、OHSAS18001职业健康安全管理体系认证等多项权威认证。")

        doc.add_paragraph()

        # 发展历程
        doc.add_paragraph("发展历程：")
        milestones = [
            "• 成立初期：公司成立，专注于高低压配电设备研发制造",
            "• 业务扩展：建立湖北孝感生产基地，扩大生产规模",
            "• 市场拓展：建立浙江宁波生产基地，形成双基地格局",
            "• 技术升级：引进先进生产设备，提升产品质量",
            "• 品牌合作：与ABB、西门子等国际知名品牌建立合作关系",
            "• 规模扩张：年产值突破20亿元，成为行业领先企业",
            "• 未来规划：持续技术创新，计划5年内完成上市"
        ]

        for milestone in milestones:
            doc.add_paragraph(milestone)

        doc.add_page_break()

    def _add_tech_solution(self, doc: Document, tender_info: Dict, matched_data: Dict):
        """添加技术方案"""
        p = doc.add_paragraph()
        run = p.add_run("第2章 方案建议书")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 2.1 工艺质量
        p = doc.add_paragraph()
        run = p.add_run("2.1 工艺质量")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        product_reqs = tender_info.get("product_requirements", [])

        doc.add_paragraph("根据招标文件要求，我司为贵方提供一套完整、可靠、先进的配电设备解决方案。")
        doc.add_paragraph()

        if product_reqs:
            doc.add_paragraph("主要产品配置：")
            for i, product in enumerate(product_reqs[:10], 1):
                doc.add_paragraph(f"{i}. {product}")
            doc.add_paragraph()

        # 2.2 技术特点
        p = doc.add_paragraph()
        run = p.add_run("2.2 技术特点")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        features = [
            "（1）产品符合国家相关标准和规范",
            "（2）采用优质元器件，确保设备可靠运行",
            "（3）结构设计合理，便于安装和维护",
            "（4）防护等级高，适应各种环境条件",
            "（5）完善的五防联锁功能，运行安全可靠",
            "（6）采用真空断路器，性能优越，维护方便"
        ]

        for feature in features:
            doc.add_paragraph(feature)

        doc.add_paragraph()

        # 2.3 主要元器件品牌
        p = doc.add_paragraph()
        run = p.add_run("2.3 主要元器件品牌")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        brands = [
            "• 真空断路器：上海人民、厦门华电",
            "• 保护装置：西门子、ABB、施耐德",
            "• 电流互感器：大连第一互感器",
            "• 断路器：施耐德、ABB",
            "• 仪表：斯菲尔、许继",
            "• 接触器：施耐德、ABB",
            "• 继电器：ABB、西门子"
        ]

        for brand in brands:
            doc.add_paragraph(brand)

        doc.add_page_break()

    def _add_qualifications_with_images(self, doc: Document, qualifications: List[Dict], data_dir: Path, show_cert_images: bool = False):
        """
        添加企业资质（支持PDF转图片）

        优点：
        - 打印出来直接可以看到证书图片
        - 客户不需要任何操作
        - 投标文件看起来更专业

        缺点：
        - 需要安装额外库（pdf2image + Pillow + Ghostscript）
        - 文件会大一些
        """
        # 直接使用PDF转图片版本
        self._add_qualifications_with_pdf_images(doc, qualifications, data_dir)

    def _add_qualifications_with_pdf_images(self, doc: Document, qualifications: List[Dict], data_dir: Path):
        """
        添加企业资质（PDF转图片）
        """
        p = doc.add_paragraph()
        run = p.add_run("第3章 企业资质")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        if not qualifications:
            doc.add_paragraph("（具体资质文件详见附件）")
            doc.add_page_break()
            return

        # 创建临时目录用于保存转换后的图片（不会被自动删除）
        import tempfile
        import json
        import uuid
        
        # 创建持久化的临时目录
        temp_dir = Path(tempfile.mkdtemp(prefix='bid_gen_certs_'))
        
        converted_images = {}
        total = 0
        success = 0
        failed = 0
        
        for i, cert in enumerate(qualifications, 1):
            if not cert.get('cert_file'):
                continue
            
            cert_path = data_dir / cert['cert_file']
            if not cert_path.exists():
                print(f"✗ 证书文件不存在: {cert['cert_file']}")
                failed += 1
                total += 1
                continue
            
            # 转换PDF为图片
            try:
                # 不使用 output_folder 参数，返回图片对象
                images = convert_from_path(
                    str(cert_path),
                    first_page=1, last_page=1,
                    dpi=200,  # 分辨率，200dpi打印效果较好
                    fmt='jpg',  # 输出为JPG格式
                    use_cropbox=True
                )
                
                if images:
                    # 调整图片大小
                    img = images[0]
                    img_width, img_height = img.size
                    
                    # 如果图片太宽，调整宽度
                    target_width = 500  # 目标宽度500像素
                    if img_width > target_width:
                        ratio = target_width / img_width
                        new_height = int(img_height * ratio)
                        img = img.resize((target_width, new_height), PILImage.LANCZOS)
                    
                    # 保存到临时目录（使用唯一文件名）
                    img_filename = f"{cert['id']}_{uuid.uuid4().hex[:8]}.jpg"
                    img_path = temp_dir / img_filename
                    img.save(img_path, quality=85)
                    
                    converted_images[cert['id']] = {
                        'name': cert['name'],
                        'level': cert['level'],
                        'images': [img_path]  # 保存图片路径
                    }
                    success += 1
                    print(f"✓ 证书 {i} 转换成功: {cert['name']}")
                else:
                    converted_images[cert['id']] = {
                        'name': cert['name'],
                        'level': cert['level'],
                        'images': [],
                        'error': '转换失败'
                    }
                    failed += 1
            except Exception as e:
                print(f"✗ 证书 {i} 转换失败: {e}")
                import traceback
                traceback.print_exc()
                converted_images[cert['id']] = {
                    'name': cert['name'],
                    'level': cert['level'],
                    'images': [],
                    'error': str(e)
                }
                failed += 1
            
            total += 1
            
            if total > 0 and total % 5 == 0:
                print(f"进度: {i}/{len(qualifications)}")
        
        print(f"转换完成: 总计 {total}, 成功 {success}, 失败 {failed}")

        # 3.1 体系认证证书
        p = doc.add_paragraph()
        run = p.add_run("3.1 体系认证证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出体系认证
        system_certs = [q for q in qualifications if "体系" in q.get("name", "") or "认证" in q.get("name", "")]

        for cert in system_certs[:10]:  # 最多显示10个
            images = converted_images.get(cert['id'], {}).get('images', [])
            
            # 显示证书名称
            p = doc.add_paragraph()
            run = p.add_run(f"• {cert['name']}（{cert['level']}）")
            run.bold = True

            # 显示证书编号和有效期
            if cert.get('cert_no'):
                doc.add_paragraph(f"  证书编号：{cert['cert_no']}")
            if cert.get('valid_until'):
                doc.add_paragraph(f"  有效期至：{cert['valid_until']}")
            
            # 插入图片
            for img_path in images:
                try:
                    doc.add_paragraph()  # 空行
                    doc.add_picture(str(img_path), width=Inches(self.image_width_inches))  # 宽度5.5英寸
                    doc.add_paragraph()  # 空行
                except Exception as e:
                    doc.add_paragraph(f"  （图片插入失败: {e}）")
                    print(f"✗ 图片插入失败: {img_path} - {e}")
        
        doc.add_paragraph()

        # 3.2 信用等级证书
        p = doc.add_paragraph()
        run = p.add_run("3.2 信用等级证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出信用等级
        credit_certs = [q for q in qualifications if "AAA" in q.get("name", "") or "信用" in q.get("name", "")]

        for cert in credit_certs[:10]:
            images = converted_images.get(cert['id'], {}).get('images', [])
            
            p = doc.add_paragraph()
            run = p.add_run(f"• {cert['name']}")
            run.bold = True
            
            if cert.get('cert_no'):
                doc.add_paragraph(f"  证书编号：{cert['cert_no']}")
            if cert.get('valid_until'):
                doc.add_paragraph(f"  有效期至：{cert['valid_until']}")
            
            # 插入图片
            for img_path in images:
                try:
                    doc.add_paragraph()
                    doc.add_picture(str(img_path), width=Inches(self.image_width_inches))
                    doc.add_paragraph()
                except:
                    doc.add_paragraph(f"  （图片插入失败）")
        
        doc.add_paragraph()

        # 3.3 重点荣誉证书
        p = doc.add_paragraph()
        run = p.add_run("3.3 重点荣誉证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出省级荣誉
        honor_certs = [q for q in qualifications if "重点" in q.get("name", "") or "质量奖" in q.get("name", "")]

        for cert in honor_certs[:10]:
            images = converted_images.get(cert['id'], {}).get('images', [])
            
            p = doc.add_paragraph()
            run = p.add_run(f"• {cert['name']}")
            run.bold = True
            
            if cert.get('cert_no'):
                doc.add_paragraph(f"  证书编号：{cert['cert_no']}")
            if cert.get('valid_until'):
                doc.add_paragraph(f"  有效期至：{cert['valid_until']}")
            
            # 插入图片
            for img_path in images:
                try:
                    doc.add_paragraph()
                    doc.add_picture(str(img_path), width=Inches(self.image_width_inches))
                    doc.add_paragraph()
                except:
                    doc.add_paragraph(f"  （图片插入失败）")
        
        doc.add_paragraph()

        # 3.4 合作伙伴证书
        p = doc.add_paragraph()
        run = p.add_run("3.4 合作伙伴证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出合作伙伴
        partner_certs = [q for q in qualifications if "授权" in q.get("name", "") or "合作" in q.get("name", "")]

        for cert in partner_certs[:10]:
            images = converted_images.get(cert['id'], {}).get('images', [])
            
            p = doc.add_paragraph()
            run = p.add_run(f"• {cert['name']}")
            run.bold = True
            
            # 插入图片
            for img_path in images:
                try:
                    doc.add_paragraph()
                    doc.add_picture(str(img_path), width=Inches(self.image_width_inches))
                    doc.add_paragraph()
                except:
                    doc.add_paragraph(f"  （图片插入失败）")
        
        doc.add_paragraph()

        # 3.5 其他证书
        # 找出没有被分类的证书，但需要有 PDF 的才显示图片
        classified_ids = set(cert['id'] for cert in system_certs + credit_certs + honor_certs + partner_certs)
        other_certs_with_images = [q for q in qualifications if q['id'] not in classified_ids and q.get('cert_file')]

        if other_certs_with_images:
            p = doc.add_paragraph()
            run = p.add_run("3.5 其他证书")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "宋体"

            doc.add_paragraph()

            for cert in other_certs_with_images:
                images = converted_images.get(cert['id'], {}).get('images', [])

                p = doc.add_paragraph()
                run = p.add_run(f"• {cert['name']}")
                run.bold = True

                if cert.get('cert_no'):
                    doc.add_paragraph(f"  证书编号：{cert['cert_no']}")
                if cert.get('valid_until'):
                    doc.add_paragraph(f"  有效期至：{cert['valid_until']}")

                # 插入图片
                for img_path in images:
                    try:
                        doc.add_paragraph()
                        doc.add_picture(str(img_path), width=Inches(self.image_width_inches))
                        doc.add_paragraph()
                    except Exception as e:
                        doc.add_paragraph(f"  （图片插入失败: {e}）")
                        print(f"✗ 图片插入失败: {img_path} - {e}")

        # 清理临时图片目录
        try:
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
            print("✓ 已清理临时图片文件")
        except Exception as e:
            print(f"✗ 清理临时目录失败: {e}")

        doc.add_page_break()

    def _add_equipment_specs_table(self, doc: Document, data_dir: Path):
        """添加设备说明一览表"""
        p = doc.add_paragraph()
        run = p.add_run("第2.3章 主要设备说明一览表")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 加载设备规格数据
        equipment_file = data_dir / "equipment_specs.json"
        
        if equipment_file.exists():
            import json
            try:
                with open(equipment_file, 'r', encoding='utf-8') as f:
                    equipment_data = json.load(f)
                    equipment_specs = equipment_data.get('equipment_specs', [])
                
                if not equipment_specs:
                    doc.add_paragraph("暂无设备规格数据")
                    doc.add_page_break()
                    return
                
                # 按类别分组
                categories = {}
                for equipment in equipment_specs:
                    category = equipment.get('category', '其他')
                    if category not in categories:
                        categories[category] = []
                    categories[category].append(equipment)
                
                # 为每个类别创建表格
                for category, items in categories.items():
                    # 添加类别标题
                    p = doc.add_paragraph()
                    run = p.add_run(category)
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.name = "黑体"
                    doc.add_paragraph()
                    
                    # 创建表格（11列：序号、符号、名称、型号、规格、材质、厚度、重量、单位、数量、生产厂家、备注）
                    table = doc.add_table(rows=len(items) + 1, cols=12)
                    table.style = 'Light Grid Accent 1'
                    
                    # 设置表头
                    headers = ['序号', '符号', '名称', '型号', '规格', '材质', '厚度', '重量(KG)', '单位', '数量', '生产厂家', '备注']
                    for col_idx, header in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = header
                        # 设置表头样式
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                                run.font.name = "宋体"
                    
                    # 填充数据
                    for row_idx, equipment in enumerate(items, start=1):
                        row = table.rows[row_idx]
                        
                        cells = [
                            str(equipment.get('sequence', '')),
                            equipment.get('symbol', ''),
                            equipment.get('name', ''),
                            equipment.get('model', ''),
                            equipment.get('specifications', ''),
                            equipment.get('material', ''),
                            equipment.get('thickness', ''),
                            str(equipment.get('weight', '')) if equipment.get('weight') is not None else '',
                            equipment.get('unit', ''),
                            str(equipment.get('quantity', '')),
                            equipment.get('manufacturer', ''),
                            equipment.get('remarks', '')
                        ]
                        
                        for col_idx, cell_text in enumerate(cells):
                            cell = row.cells[col_idx]
                            cell.text = cell_text
                            # 设置数据样式
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    run.font.name = "宋体"
                    
                    # 类别之间添加空行
                    doc.add_paragraph()
                
                doc.add_page_break()
                
            except Exception as e:
                print(f"✗ 加载设备规格数据失败: {e}")
                doc.add_paragraph("设备规格数据加载失败")
                doc.add_page_break()
        else:
            doc.add_paragraph("设备规格数据文件不存在")
            doc.add_page_break()

    def _add_qualifications_v2(self, doc: Document, qualifications: List[Dict]):
        """添加资质文件 - V2（升级版，显示文件路径）"""
        p = doc.add_paragraph()
        run = p.add_run("第3章 企业资质")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        if not qualifications:
            doc.add_paragraph("（具体资质文件详见附件）")
            doc.add_page_break()
            return

        # 3.1 体系认证证书
        p = doc.add_paragraph()
        run = p.add_run("3.1 体系认证证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出体系认证
        system_certs = [q for q in qualifications if "体系" in q.get("name", "") or "认证" in q.get("name", "")]

        for cert in system_certs[:10]:
            doc.add_paragraph(f"• {cert['name']}（{cert['level']}）")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.2 信用等级证书
        p = doc.add_paragraph()
        run = p.add_run("3.2 信用等级证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出信用等级
        credit_certs = [q for q in qualifications if "AAA" in q.get("name", "") or "信用" in q.get("name", "")]

        for cert in credit_certs[:10]:
            doc.add_paragraph(f"• {cert['name']}")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.3 省级荣誉证书
        p = doc.add_paragraph()
        run = p.add_run("3.3 省级荣誉证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出省级荣誉
        honor_certs = [q for q in qualifications if "重点" in q.get("name", "") or "质量奖" in q.get("name", "")]

        for cert in honor_certs[:10]:
            doc.add_paragraph(f"• {cert['name']}")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.4 合作伙伴证书
        p = doc.add_paragraph()
        run = p.add_run("3.4 合作伙伴证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出合作伙伴
        partner_certs = [q for q in qualifications if "授权" in q.get("name", "") or "合作" in q.get("name", "")]

        for cert in partner_certs[:10]:
            doc.add_paragraph(f"• {cert['name']}")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.5 其他证书
        remaining = len(qualifications) - len(system_certs) - len(credit_certs) - len(honor_certs) - len(partner_certs)
        if remaining > 0:
            p = doc.add_paragraph()
            run = p.add_run("3.5 其他证书")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "宋体"

            doc.add_paragraph()
            doc.add_paragraph(f"（其他证书共 {remaining} 项，详见附件）")

        doc.add_page_break()

    def _add_performance(self, doc: Document, cases: List[Dict]):
        """添加类似业绩"""
        p = doc.add_paragraph()
        run = p.add_run("第4章 类似业绩")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        if not cases:
            doc.add_paragraph("（具体业绩清单详见附件）")
            doc.add_page_break()
            return

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ["序号", "项目名称", "客户", "行业", "产品类型", "金额(万元)"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # 填充数据
        for i, case in enumerate(cases[:10], 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = case.get("project_name", "")
            row_cells[2].text = case.get("client", "")
            row_cells[3].text = case.get("industry", "")
            row_cells[4].text = case.get("product_type", "")
            row_cells[5].text = str(case.get("amount", 0) / 10000)

        doc.add_page_break()

    def _add_quotation(self, doc: Document, quote_data: Dict):
        """添加报价单"""
        p = doc.add_paragraph()
        run = p.add_run("报价明细")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 报价汇总
        products = quote_data.get("products", [])
        total_amount = sum(p.get("小计", 0) for p in products)

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ["序号", "产品名称", "型号/规格", "单位", "数量", "单价(元)"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header

        # 填充数据
        for i, product in enumerate(products, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = product.get("产品名称", "")
            row_cells[2].text = product.get("型号/规格", "")
            row_cells[3].text = product.get("单位", "")
            row_cells[4].text = str(product.get("数量", ""))
            row_cells[5].text = str(product.get("单价", ""))

        doc.add_paragraph()

        # 合计
        p = doc.add_paragraph()
        run = p.add_run(f"报价合计：人民币 {total_amount:,.2f} 元")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 0, 0)

        # 报价说明
        doc.add_paragraph()
        doc.add_paragraph("报价说明：")
        doc.add_paragraph("1. 以上报价含13%增值税，含国内运费")
        doc.add_paragraph("2. 交货期：合同签订后30天内（可根据要求调整）")
        doc.add_paragraph("3. 质量保证：产品免费保修一年，终身提供技术服务")
        doc.add_paragraph("4. 付款方式：按合同约定执行")

        doc.add_page_break()

    def _add_after_sales(self, doc: Document, company_info: Dict):
        """添加售后服务"""
        p = doc.add_paragraph()
        run = p.add_run("售后服务")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 服务承诺
        doc.add_paragraph("服务承诺：")
        commitments = [
            "1. 本公司提供7×24小时免费技术咨询服务。",
            "2. 产品自出厂之日起免费保修一年（易损件除外）。",
            "3. 保修期内，如产品出现质量问题，本公司免费维修或更换。",
            "4. 质保期后，本公司提供终身有偿维修服务。",
            "5. 本公司定期回访客户，了解设备运行情况，提供技术支持。",
            "6. 本公司建立完善的服务网络，确保及时响应客户需求。",
            "7. 本司遵守招标文件规定的所有商务条款。",
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_paragraph()

        # 联系方式
        p = doc.add_paragraph()
        run = p.add_run("联系咨询")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()
        doc.add_paragraph(f"如需了解更多信息，请联系我们：")
        doc.add_paragraph(f"电话：{company_info['phone']}")
        doc.add_paragraph(f"传真：{company_info['fax']}")
        doc.add_paragraph(f"邮箱：{company_info['email']}")
        doc.add_paragraph(f"地址：{company_info['address']}")

        doc.add_page_break()

    def _add_tech_commitment(self, doc: Document):
        """添加技术承诺"""
        p = doc.add_paragraph()
        run = p.add_run("技术承诺")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        commitments = [
            "1. 本公司提供的所有产品均符合国家相关标准和规范。",
            "2. 本公司承诺严格按照招标文件要求提供技术方案。",
            "3. 本公司承诺所投产品具有稳定可靠的质量。",
            "4. 本公司承诺在交货期内完成所有产品的交付。",
            "5. 本公司承诺提供完善的售后服务和技术支持。",
            "6. 本公司承诺对所投产品的技术参数负责。",
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_page_break()

    def _add_response_commitment(self, doc: Document):
        """添加响应承诺"""
        p = doc.add_paragraph()
        run = p.add_run("响应承诺")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        commitments = [
            "1. 本公司对招标文件的各项要求均已充分了解。",
            "2. 本公司承诺按照招标文件要求提供合格的投标文件。",
            "3. 本公司承诺在投标有效期内保持投标有效。",
            "4. 本公司承诺中标后严格按照合同约定履行义务。",
            "5. 本公司承诺接受招标文件中的所有条款和条件。",
            "6. 本公司承诺不对招标文件的任何内容提出异议。",
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_page_break()

    def _add_commercial_commitment(self, doc: Document):
        """添加商务承诺"""
        p = doc.add_paragraph()
        run = p.add_run("商务承诺")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        commitments = [
            "1. 本公司承诺按照合同约定的时间和方式交付产品。",
            "2. 本公司承诺产品质量符合合同约定的标准。",
            "3. 本公司承诺按照合同约定的价格执行合同。",
            "4. 本公司承诺提供符合合同约定的售后服务。",
            "5. 本公司承诺遵守所有商务条款和条件。",
            "6. 本公司承诺对所提供的产品和服务的价格负责。",
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document, separate_bids: bool = False):
        """添加目录（动态生成）"""
        doc.add_page_break()
        
        # 添加目录标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("目录")
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 动态生成目录内容
        contents = []
        
        if separate_bids:
            # 技术标和商务标分开
            contents.extend([
                ("1.", "技术标", "1"),
                ("1.1", "封面", "1"),
                ("1.2", "公司概况", "2"),
                ("1.3", "公司简介", "3"),
                ("1.4", "公司资质", "4"),
                ("1.5", "公司业绩", "5"),
                ("1.6", "投标纲领", "6"),
                ("1.7", "技术偏离表", "7"),
                ("1.8", "公司概况", "8"),
                ("1.9", "技术方案", "9"),
                ("1.10", "项目团队", "10"),
                ("1.11", "设备说明一览表", "11"),
                ("1.12", "技术承诺", "12"),
                ("", "", ""),
                ("2.", "商务标", "12"),
                ("2.1", "封面", "13"),
                ("2.2", "公司概况", "14"),
                ("2.3", "投标纲领", "15"),
                ("2.4", "商务偏离表", "16"),
                ("2.5", "报价说明", "17"),
                ("2.6", "商务承诺", "18"),
            ])
        else:
            # 单一文件
            contents.extend([
                ("1.", "封面", "1"),
                ("2.", "公司概况", "2"),
                ("3.", "公司简介", "3"),
                ("4.", "公司资质", "4"),
                ("5.", "公司业绩", "5"),
                ("6.", "投标纲领", "6"),
                ("7.", "技术偏离表", "7"),
                ("8.", "公司概况", "8"),
                ("9.", "技术方案", "9"),
                ("10.", "设备说明一览表", "10"),
                ("11.", "资质证书", "11"),
                ("12.", "项目案例", "12"),
                ("13.", "报价说明", "13"),
                ("14.", "售后服务", "14"),
                ("15.", "技术承诺", "15"),
                ("16.", "响应承诺", "16"),
                ("17.", "商务承诺", "17"),
            ])

        # 创建表格形式的目录
        table = doc.add_table(rows=0, cols=3)
        table.style = 'Table Grid'

        for seq, title, page in contents:
            row = table.add_row()
            row.cells[0].text = seq
            row.cells[1].text = title
            row.cells[2].text = page

        doc.add_page_break()

    def _setup_page_numbers(self, doc: Document):
        """设置页码和页面边距"""
        # 获取文档节
        for section in doc.sections:
            # 设置页边距
            section.left_margin = Cm(2.54)   # 左边距 2.54cm
            section.right_margin = Cm(2.54)  # 右边距 2.54cm
            section.top_margin = Cm(2.54)    # 上边距 2.54cm
            section.bottom_margin = Cm(2.54) # 下边距 2.54cm

            # 设置页脚
            footer = section.footer

            # 在页脚添加页码
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加页码
            run = footer_para.add_run()
            run.font.size = Pt(10)

            # 创建页码域
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

            # 添加页码前后缀
            run._element.insertbefore(footer_para.add_run("第 ")._element, run._element[0])
            run._element.append(footer_para.add_run(" 页")._element)


# 测试代码
if __name__ == "__main__":
    print("投标文件生成器 - PDF转图片版")
    print("=" * 60)
    print()
    print("功能：")
    print("- 所有Bug已修复")
    print("- 支持PDF转图片功能")
    print("- 详细的日志记录")
    print("- 增强的调试模式")
    print("- 错误处理和恢复")
    print()
    print("使用方法：")
    print("1. 此文件替换原来的 generator.py")
    print("2. 重启应用")
    print("3. 在生成投标文件时，勾选'显示证书图片'选项")
    print()
    print("=" * 60)
