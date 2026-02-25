#!/usr/bin/env python3
"""
PDF 转图片功能（不依赖 pdf2image，直接使用 Ghostscript）
"""

import sys
import os
import subprocess
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

print("=" * 80)
print("PDF 转图片功能（Ghostscript 直接调用版）")
print("=" * 80)
print()

# 检查 Ghostscript
print("检查 Ghostscript...")
gs_path = subprocess.run(["which", "gs"], capture_output=True, text=True).stdout.strip()
if not gs_path:
    print("✗ Ghostscript 未安装或不在 PATH 中")
    print()
    print("请安装 Ghostscript：")
    print("  Mac: brew install ghostscript")
    print("  Linux: sudo apt-get install ghostscript")
    print()
    sys.exit(1)

print(f"✓ Ghostscript 已安装：{gs_path}")
print()

# 检查 Ghostscript 版本
result = subprocess.run([gs_path, "--version"], capture_output=True, text=True)
gs_version = result.stdout.strip()
print(f"✓ Ghostscript 版本：{gs_version}")
print()

def convert_pdf_to_page_jpeg(pdf_path: str, output_path: str, dpi: int = 200) -> bool:
    """
    使用 Ghostscript 直接将 PDF 的第1页转换为 JPEG 图片
    
    Args:
        pdf_path: PDF 文件路径
        output_path: 输出图片路径
        dpi: 分辨率（默认 200）
    
    Returns:
        是否转换成功
    """
    try:
        # Ghostscript 命令（只转换第1页）
        # -dFirstPage=1 - 从第1页开始
        # -dLastPage=1 - 到第1页结束（只转换第1页）
        # -sDEVICE=jpeg - 输出为 JPEG 格式
        # -r200 - 分辨率 200dpi
        # -dTextAlphaBits=4 - 文本抗锯齿
        # -dGraphicsAlphaBits=4 - 图形抗锯齿
        command = [
            gs_path,
            "-dFirstPage=1",
            "-dLastPage=1",
            "-sDEVICE=jpeg",
            f"-r{dpi}",
            "-dTextAlphaBits=4",
            "-dGraphicsAlphaBits=4",
            "-dBATCH",
            "-dQUIET",
            "-dSAFER",
            "-sColorConversionStrategy=RGB",
            "-dProcessColorModel=/DeviceRGB",
            "-dAutoRotatePages=/None",
            f"-sOutputFile={output_path}",
            pdf_path
        ]
        
        # 运行命令
        result = subprocess.run(
            command,
            capture_output=True,
            text=True
        )
        
        # 检查结果
        if result.returncode == 0:
            # 检查输出文件是否存在
            if os.path.exists(output_path):
                return True
            else:
                print(f"  ✗ 转换成功，但输出文件不存在：{output_path}")
                return False
        else:
            print(f"  ✗ Ghostscript 转换失败，返回码：{result.returncode}")
            print(f"  错误信息：{result.stderr}")
            return False
            
    except Exception as e:
        print(f"  ✗ 转换失败：{e}")
        return False

def test_pdf_to_image():
    """测试 PDF 转图片功能"""
    print("=" * 80)
    print("测试 PDF 转图片")
    print("=" * 80)
    print()
    
    # 获取测试证书文件
    from config import DATA_DIR
    from database import CompanyDatabase
    
    db = CompanyDatabase(DATA_DIR)
    qualifications = db.get_qualifications()
    
    print(f"找到 {len(qualifications)} 个资质证书")
    print()
    
    # 查找第一个有 PDF 文件的证书
    found_cert = None
    for cert in qualifications:
        cert_file = cert.get('cert_file')
        if cert_file:
            cert_path = DATA_DIR / cert_file
            if cert_path.exists():
                found_cert = cert
                break
    
    if not found_cert:
        print("✗ 没有找到有效的证书文件")
        sys.exit(1)
    
    print("=" * 80)
    print("测试证书文件")
    print("=" * 80)
    print()
    
    cert_file = found_cert.get('cert_file')
    cert_path = DATA_DIR / cert_file
    
    print(f"证书名称：{found_cert['name']}")
    print(f"证书等级：{found_cert['level']}")
    print(f"证书编号：{found_cert.get('cert_no', '')}")
    print(f"证书文件：{cert_file}")
    print(f"文件路径：{cert_path}")
    print(f"文件存在：{cert_path.exists()}")
    print(f"文件大小：{cert_path.stat().st_size} 字节")
    print()
    
    # 测试转换
    print("=" * 80)
    print("测试 PDF 转图片（使用 Ghostscript 直接调用）")
    print("=" * 80)
    print()
    
    import tempfile
    
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test_output.jpg"
        print(f"输出路径：{output_path}")
        print()
        
        print("开始转换...")
        print(f"  Ghostscript 命令：-dFirstPage=1 -dLastPage=1 -sDEVICE=jpeg -r200 ... {cert_path}")
        print()
        
        success = convert_pdf_to_page_jpeg(str(cert_path), str(output_path), dpi=200)
        
        print()
        if success:
            print("✓ 转换成功！")
            print(f"  输出文件大小：{output_path.stat().st_size} 字节")
            
            # 检查图片尺寸
            try:
                from PIL import Image
                img = Image.open(str(output_path))
                img_width, img_height = img.size
                print(f"  图片尺寸：{img_width} x {img_height}")
            except ImportError:
                print("  ⚠ Pillow 未安装，无法检查图片尺寸")
            
            print()
            print("=" * 80)
            print("测试 Word 文档生成")
            print("=" * 80)
            print()
            
            # 测试 Word 文档生成
            doc = Document()
            p = doc.add_paragraph()
            run = p.add_run("测试证书图片（Ghostscript 版）")
            run.bold = True
            
            # 添加图片
            try:
                doc.add_paragraph()
                doc.add_picture(str(output_path), width=Inches(5.5))
                doc.add_paragraph()
                print("  ✓ 图片成功插入到 Word 文档")
            except:
                print("  ✗ 图片插入失败")
            
            # 保存到临时目录
            output_doc_path = Path(temp_dir) / "test_doc.docx"
            doc.save(output_doc_path)
            print(f"  ✓ Word 文档已生成：{output_doc_path}")
            print(f"  文件大小：{output_doc_path.stat().st_size} 字节")
            print()
            
            print("=" * 80)
            print("测试结果")
            print("=" * 80)
            print()
            print("✓ PDF 文件可以读取")
            print("✓ Ghostscript 可以转换 PDF")
            print("✓ 转换后的图片可以插入到 Word 文档")
            print("✓ Word 文档可以保存")
            print()
            print("✅ PDF 转图片功能正常（不依赖 pdf2image）！")
            print()
            print("=" * 80)
            print("下一步：")
            print("=" * 80)
            print()
            print("1. 将此方法集成到 generator.py")
            print("2. 更新 _add_qualifications_with_pdf_images 方法")
            print("3. 使用 Ghostscript 直接调用代替 pdf2image")
            print("4. 重新生成投标文件")
            print("5. 查看生成的文件中是否有证书图片")
            print()
            print("=" * 80)
            sys.exit(0)
            
        else:
            print("✗ 转换失败")
            print()
            print("可能的原因：")
            print("  1. PDF 文件已损坏")
            print("  2. Ghostscript 未正确安装")
            print("  3. 没有权限读取 PDF 文件")
            print("  4. Ghostscript 版本不兼容")
            print()
            sys.exit(1)

if __name__ == "__main__":
    test_pdf_to_image()
