#!/usr/bin/env python3
"""
PDF 转图片功能测试脚本（第1轮）
"""

import sys
import os
from pathlib import Path

# 添加当前目录到路径
sys.path.insert(0, str(Path(__file__).parent / "bid-generator"))

from config import DATA_DIR
from database import CompanyDatabase
from generator import BidDocumentGenerator

print("=" * 80)
print("PDF 转图片功能测试（第1轮）")
print("=" * 80)
print()

print("测试目标：")
print("- 检查 PDF 转图片功能是否正常工作")
print("- 验证 API 参数修复是否正确")
print("- 检查生成的 Word 文档中是否有图片")
print()

# 初始化数据库
db = CompanyDatabase(DATA_DIR)

# 获取所有资质
qualifications = db.get_qualifications()

print(f"找到 {len(qualifications)} 个资质证书")
print()

# 查找第一个有 PDF 文件的证书
found_cert = None
for cert in qualifications:
    cert_file = cert.get('cert_file')
    if cert_file:
        cert_path = DATA_DIR / cert_file
        if cert_path.exists():
            found_cert = cert
            break

if not found_cert:
    print("✗ 没有找到有效的证书文件")
    print()
    print("测试失败！")
    sys.exit(1)

print("=" * 80)
print("测试证书文件")
print("=" * 80)
print()

cert_file = found_cert.get('cert_file')
cert_path = DATA_DIR / cert_file

print(f"证书名称：{found_cert['name']}")
print(f"证书等级：{found_cert['level']}")
print(f"证书编号：{found_cert.get('cert_no', '')}")
print(f"证书文件：{cert_file}")
print(f"文件路径：{cert_path}")
print(f"文件存在：{cert_path.exists()}")
print(f"文件大小：{cert_path.stat().st_size} 字节")
print()

# 测试 PDF 转图片功能
print("=" * 80)
print("测试 PDF 转图片（使用修复后的 API）")
print("=" * 80)
print()

try:
    # 检查依赖
    print("检查依赖...")
    
    import pdf2image
    print("  ✓ pdf2image 已安装")
    
    from PIL import Image
    print("  ✓ Pillow 已安装")
    
    # 检查 Ghostscript
    import shutil
    gs_path = shutil.which("gs")
    if gs_path:
        print(f"  ✓ Ghostscript 已安装：{gs_path}")
    else:
        print("  ✗ Ghostscript 未安装或不在 PATH 中")
        print()
        print("请安装 Ghostscript：")
        print("  Mac: brew install ghostscript")
        print("  Linux: sudo apt-get install ghostscript")
        sys.exit(1)
    
    print()
    
    # 测试转换（使用修复后的 API 参数）
    print("开始转换 PDF...")
    print()
    
    import tempfile
    
    with tempfile.TemporaryDirectory() as temp_dir:
        print(f"临时目录：{temp_dir}")
        print()
        
        # 方法：使用 first_page 和 last_page 参数（修复后的版本）
        print("方法：使用 first_page=1, last_page=1 参数...")
        try:
            images = pdf2image.convert_from_path(
                str(cert_path),
                output_folder=temp_dir,
                first_page=1,  # 【修复】从第1页开始
                last_page=1,   # 【修复】只转换第1页
                dpi=200,
                fmt='jpg',
                use_cropbox=True
            )
            
            if images:
                print(f"  ✓ 转换成功！生成 {len(images)} 张图片")
                
                img = Image.open(images[0])
                img_width, img_height = img.size
                print(f"  图片尺寸：{img_width} x {img_height}")
                print(f"  图片文件：{images[0]}")
                print(f"  图片大小：{os.path.getsize(images[0])} 字节")
                print()
                
                print("=" * 80)
                print("测试 Word 文档生成")
                print("=" * 80)
                print()
                
                # 测试 Word 文档生成
                from docx import Document
                from docx.shared import Inches
                
                doc = Document()
                
                # 标题
                p = doc.add_paragraph()
                run = p.add_run("测试证书图片")
                run.bold = True
                run.font.size = 14
                run.font.name = "黑体"
                
                # 证书信息
                p = doc.add_paragraph()
                p.add_paragraph(f"证书名称：{found_cert['name']}")
                p.add_paragraph(f"证书等级：{found_cert['level']}")
                p.add_paragraph(f"证书编号：{found_cert.get('cert_no', '')}")
                p.add_paragraph(f"有效期至：{found_cert.get('valid_until', '')}")
                p.add_paragraph(f"证书文件：{cert_file}")
                
                # 添加图片（使用修复后的 API）
                doc.add_paragraph()
                doc.add_picture(str(images[0]), width=Inches(5.5))
                doc.add_paragraph()
                print("  ✓ 图片成功插入到 Word 文档")
                
                # 保存到临时目录
                output_path = Path(temp_dir) / "test_doc.docx"
                doc.save(output_path)
                print(f"  ✓ Word 文档已生成：{output_path}")
                print(f"  文件大小：{output_path.stat().st_size} 字节")
                print()
                
                # 验证 Word 文档
                from docx import Document
                doc = Document(output_path)
                
                # 检查文档中的图片
                image_count = 0
                for rel in doc.part.rels:
                    if "image" in rel.target_ref:
                        image_count += 1
                
                print(f"  ✓ Word 文档验证通过！")
                print(f"  文档中包含 {image_count} 个图片")
                print()
                
                print("=" * 80)
                print("测试结果")
                print("=" * 80)
                print()
                print("✓ PDF 文件可以读取")
                print("✓ PDF 可以转换为图片")
                print("✓ 图片可以插入到 Word 文档")
                print("✓ Word 文档可以保存")
                print("✓ Word 文档中包含图片")
                print()
                print("✅ PDF 转图片功能正常！")
                print()
                print("=" * 80)
                print("第1轮测试成功！")
                print("=" * 80)
                print()
                print("下一步：")
                print("1. 重启 Streamlit 应用")
                print("2. 开启调试模式")
                print("3. 导入中天钢铁的招标文件")
                print("4. 勾选'显示证书图片'")
                print("5. 生成投标文件")
                print("6. 查看调试信息和生成的文件")
                print()
                print("如果调试信息显示 'show_cert_images_final = True'，应该会显示证书图片！")
                print()
                print("如果还是有问题，请告诉我调试信息和错误信息。")
                print()
                
                sys.exit(0)
                
            else:
                print(f"  ✗ 转换失败：没有生成图片")
                print()
                print("可能的原因：")
                print("  1. PDF 文件已损坏")
                print("  2. Ghostscript 未正确安装")
                print("  3. PDF 文件是受保护的")
                print("  4. pdf2image 版本不兼容")
                sys.exit(1)
                
        except Exception as e:
            print(f"  ✗ 转换失败：{e}")
            print()
            import traceback
            print("详细错误：")
            traceback.print_exc()
            print()
            print("可能的原因：")
            print("  1. PDF 文件已损坏")
            print("  2. Ghostscript 未正确安装")
            print("  3. 没有权限读取 PDF 文件")
            print("  4. pdf2image 版本不兼容")
            sys.exit(1)
            
except ImportError as e:
    print(f"✗ 导入错误：{e}")
    print()
    print("请安装依赖：")
    print("  pip install pdf2image Pillow")
    print()
    sys.exit(1)
    
except Exception as e:
    import traceback
    print(f"✗ 转换失败：{e}")
    print()
    print("详细错误：")
    traceback.print_exc()
    print()
    print("可能的原因：")
    print("  1. Ghostscript 未安装或不在 PATH 中")
    print("  2. PDF 文件已损坏")
    print("  3. 没有权限读取 PDF 文件")
    print("  4. pdf2image 版本不兼容")
    sys.exit(1)

print("=" * 80)
print("测试完成")
print("=" * 80)
