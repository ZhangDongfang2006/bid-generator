#!/usr/bin/env python3
"""
测试实际的生成流程，模拟 BidDocumentGenerator 的调用
"""

import sys
sys.path.insert(0, '/Users/zhangdongfang/workspace/bid-generator')

from pathlib import Path
import json
from docx import Document
from docx.shared import Pt, Inches
from pdf2image import convert_from_path
from PIL import Image as PILImage
import tempfile
import uuid

# 读取资质数据
data_dir = Path('/Users/zhangdongfang/workspace/bid-generator/data')
with open(data_dir / 'qualifications.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

qualifications = data.get('qualifications', [])

print(f"总共有 {len(qualifications)} 个证书")
print()

# 按分类统计
system_certs = [q for q in qualifications if "体系" in q.get("name", "") or "认证" in q.get("name", "")]
credit_certs = [q for q in qualifications if "AAA" in q.get("name", "") or "信用" in q.get("name", "")]
honor_certs = [q for q in qualifications if "重点" in q.get("name", "") or "质量奖" in q.get("name", "")]
partner_certs = [q for q in qualifications if "授权" in q.get("name", "") or "合作" in q.get("name", "")]

print(f"3.1 体系认证证书: {len(system_certs)} 个")
for cert in system_certs:
    has_file = '有PDF' if cert.get('cert_file') else '无PDF'
    print(f"  - {cert['name']} [{has_file}]")
print()

print(f"3.2 信用等级证书: {len(credit_certs)} 个")
print(f"3.3 重点荣誉证书: {len(honor_certs)} 个")
print(f"3.4 合作伙伴证书: {len(partner_certs)} 个")
print()

# 转换所有有 PDF 的证书
temp_dir = Path(tempfile.mkdtemp(prefix='test_real_flow_'))
print(f"临时目录: {temp_dir}")
print()

converted_images = {}
total = 0
success = 0
failed = 0

for i, cert in enumerate(qualifications, 1):
    if not cert.get('cert_file'):
        continue
    
    cert_path = data_dir / cert['cert_file']
    if not cert_path.exists():
        print(f"✗ 证书文件不存在: {cert['cert_file']}")
        failed += 1
        total += 1
        continue
    
    # 转换PDF为图片
    try:
        images = convert_from_path(
            str(cert_path),
            first_page=1,
            last_page=1,
            dpi=200,
            fmt='jpg',
            use_cropbox=True
        )
        
        if images:
            img = images[0]
            img_width, img_height = img.size
            
            target_width = 500
            if img_width > target_width:
                ratio = target_width / img_width
                new_height = int(img_height * ratio)
                img = img.resize((target_width, new_height), PILImage.LANCZOS)
            
            img_filename = f"{cert['id']}_{uuid.uuid4().hex[:8]}.jpg"
            img_path = temp_dir / img_filename
            img.save(img_path, quality=85)
            
            converted_images[cert['id']] = {
                'name': cert['name'],
                'level': cert['level'],
                'images': [img_path]
            }
            success += 1
            print(f"✓ 证书 {i} 转换成功: {cert['name']}")
        else:
            converted_images[cert['id']] = {
                'name': cert['name'],
                'level': cert['level'],
                'images': [],
                'error': '转换失败'
            }
            failed += 1
    except Exception as e:
        print(f"✗ 证书 {i} 转换失败: {e}")
        import traceback
        traceback.print_exc()
        converted_images[cert['id']] = {
            'name': cert['name'],
            'level': cert['level'],
            'images': [],
            'error': str(e)
        }
        failed += 1
    
    total += 1
    
    if total > 0 and total % 5 == 0:
        print(f"进度: {i}/{len(qualifications)}")

print()
print(f"转换完成: 总计 {total}, 成功 {success}, 失败 {failed}")
print()

# 创建Word文档
doc = Document()

# 3.1 体系认证证书
print("插入 3.1 体系认证证书...")
p = doc.add_paragraph()
run = p.add_run("3.1 体系认证证书")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

for cert in system_certs[:10]:  # 最多显示10个
    images = converted_images.get(cert['id'], {}).get('images', [])
    
    print(f"  处理证书: {cert['name']}, id={cert['id']}, 图片数={len(images)}")
    
    p = doc.add_paragraph()
    run = p.add_run(f"• {cert['name']}（{cert['level']}）")
    run.bold = True

    if cert.get('cert_no'):
        doc.add_paragraph(f"  证书编号：{cert['cert_no']}")
    if cert.get('valid_until'):
        doc.add_paragraph(f"  有效期至：{cert['valid_until']}")
    
    # 插入图片
    for img_path in images:
        try:
            doc.add_paragraph()
            doc.add_picture(str(img_path), width=Inches(5.5))
            doc.add_paragraph()
            print(f"    ✓ 图片插入成功")
        except Exception as e:
            p.add_paragraph(f"  （图片插入失败: {e}）")
            print(f"    ✗ 图片插入失败: {e}")

doc.add_paragraph()

# 3.5 其他证书
print("插入 3.5 其他证书...")
classified_ids = set(cert['id'] for cert in system_certs + credit_certs + honor_certs + partner_certs)
other_certs_with_images = [q for q in qualifications if q['id'] not in classified_ids and q.get('cert_file')]

print(f"  未分类但有PDF的证书: {len(other_certs_with_images)} 个")
for cert in other_certs_with_images:
    images = converted_images.get(cert['id'], {}).get('images', [])
    print(f"  - {cert['name']}, id={cert['id']}, 图片数={len(images)}")

if other_certs_with_images:
    p = doc.add_paragraph()
    run = p.add_run("3.5 其他证书")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    for cert in other_certs_with_images:
        images = converted_images.get(cert['id'], {}).get('images', [])

        p = doc.add_paragraph()
        run = p.add_run(f"• {cert['name']}")
        run.bold = True

        if cert.get('cert_no'):
            doc.add_paragraph(f"  证书编号：{cert['cert_no']}")
        if cert.get('valid_until'):
            doc.add_paragraph(f"  有效期至：{cert['valid_until']}")

        # 插入图片
        for img_path in images:
            try:
                doc.add_paragraph()
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.add_paragraph()
                print(f"    ✓ {cert['name']} 图片插入成功")
            except Exception as e:
                p.add_paragraph(f"  （图片插入失败: {e}）")
                print(f"    ✗ {cert['name']} 图片插入失败: {e}")

# 保存Word文档
output_path = Path('/tmp/test_real_flow.docx')
doc.save(output_path)

# 清理临时目录
import shutil
shutil.rmtree(temp_dir, ignore_errors=True)

print()
print("=" * 60)
print(f"✓ Word文档已保存: {output_path}")
print(f"✓ 已清理临时目录")
print()
print(f"统计:")
print(f"  - 总证书: {len(qualifications)} 个")
print(f"  - 体系认证: {len(system_certs)} 个")
print(f"  - 有PDF的体系认证: {len([c for c in system_certs if c.get('cert_file')])} 个")
print(f"  - 未分类有PDF: {len(other_certs_with_images)} 个")
print(f"  - 转换成功: {success} 个")
print(f"  - 转换失败: {failed} 个")
print()
print("请打开 Word 文档查看")
