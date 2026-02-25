#!/usr/bin/env python3
"""
修复 generator.py 中的 PDF 转图片功能
使用 Ghostscript 直接调用，不依赖 pdf2image 的 convert_from_path
"""

import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def convert_pdf_to_jpeg_gs(pdf_path: str, output_path: str, dpi: int = 200) -> bool:
    """
    使用 Ghostscript 直接将 PDF 的第1页转换为 JPEG 图片
    
    Args:
        pdf_path: PDF 文件路径
        output_path: 输出图片路径
        dpi: 分辨率（默认 200）
    
    Returns:
        是否转换成功
    """
    try:
        # 获取 Ghostscript 路径
        gs_path = subprocess.run(["which", "gs"], capture_output=True, text=True).stdout.strip()
        
        if not gs_path:
            print("✗ Ghostscript 未安装或不在 PATH 中")
            return False
        
        # Ghostscript 命令（只转换第1页）
        # -dFirstPage=1 - 从第1页开始
        # -sDEVICE=jpeg - 输出为 JPEG 格式
        # -r{dpi} - 分辨率
        # -dJPEGQ=95 - JPEG 质量 95
        # -dNOPAUSE - 不暂停
        # -dBATCH - 批处理
        # -dQUIET - 安静模式
        # -sPAPERSIZE=a4 - A4 纸张大小
        # -sOutputFile={output_path} - 输出文件路径
        command = [
            gs_path,
            "-dFirstPage=1",
            "-sDEVICE=jpeg",
            f"-r{dpi}",
            "-dJPEGQ=95",
            "-dNOPAUSE",
            "-dBATCH",
            "-dQUIET",
            f"-sOutputFile={output_path}",
            pdf_path
        ]
        
        # 运行命令
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=30  # 30 秒超时
        )
        
        # 检查结果
        if result.returncode == 0 and Path(output_path).exists():
            # 调整图片大小（如果需要）
            img = PILImage.open(output_path)
            img_width, img_height = img.size
            
            # 如果图片太宽，调整宽度
            target_width = 1100  # 目标宽度 1100 像素（约 14 厘米）
            if img_width > target_width:
                ratio = target_width / img_width
                new_height = int(img_height * ratio)
                img = img.resize((target_width, new_height), PILImage.LANCZOS)
                img.save(output_path, quality=85)
            
            return True
        else:
            print(f"✗ Ghostscript 转换失败，返回码：{result.returncode}")
            if result.stderr:
                print(f"  错误信息：{result.stderr}")
            return False
            
    except Exception as e:
        print(f"✗ 转换失败：{e}")
        return False

# 测试
if __name__ == "__main__":
    print("=" * 80)
    print("Ghostscript PDF 转图片修复脚本")
    print("=" * 80)
    print()
    
    # 检查 PDF 文件
    cert_path = Path("data/certificates/03、认证证书/01、质量管理体系认证证书-中英文版.pdf")
    
    if not cert_path.exists():
        print("✗ PDF 文件不存在")
        sys.exit(1)
    
    print(f"PDF 文件：{cert_path}")
    print(f"文件大小：{cert_path.stat().st_size} 字节")
    print()
    
    # 测试转换
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test_output.jpg"
        print(f"输出路径：{output_path}")
        print()
        
        print("开始转换...")
        success = convert_pdf_to_jpeg_gs(str(cert_path), str(output_path), dpi=200)
        print()
        
        if success:
            print("✓ 转换成功！")
            print(f"输出文件大小：{output_path.stat().st_size} 字节")
            print()
            
            # 检查图片
            img = PILImage.open(str(output_path))
            img_width, img_height = img.size
            print(f"图片尺寸：{img_width} x {img_height}")
            print(f"图片文件：{output_path}")
            print()
            
            # 测试 Word 文档生成
            from docx.shared import Inches
            
            doc = Document()
            p = doc.add_paragraph()
            run = p.add_run("测试证书图片")
            run.bold = True
            run.font.size = Pt(14)
            
            # 添加图片
            try:
                doc.add_paragraph()
                doc.add_picture(str(output_path), width=Inches(5.5))
                doc.add_paragraph()
                print("✓ 图片成功插入到 Word 文档")
            except Exception as e:
                print(f"✗ 图片插入失败：{e}")
            
            # 保存到临时目录
            output_doc_path = Path(temp_dir) / "test_doc.docx"
            doc.save(output_doc_path)
            print(f"✓ Word 文档已生成：{output_doc_path}")
            print(f"  文件大小：{output_doc_path.stat().st_size} 字节")
            print()
            
            print("=" * 80)
            print("测试结果")
            print("=" * 80)
            print()
            print("✓ PDF 文件可以读取")
            print("✓ PDF 可以转换为图片")
            print("✓ 图片可以插入到 Word 文档")
            print("✓ Word 文档可以保存")
            print()
            print("✅ Ghostscript 直接调用方式正常！")
            print()
            print("=" * 80)
            print("下一步：")
            print("=" * 80)
            print()
            print("1. 将此函数集成到 generator.py")
            print("2. 更新 _add_qualifications_with_pdf_images 方法")
            print("3. 使用 Ghostscript 直接调用代替 pdf2image")
            print("4. 重新生成投标文件")
            print()
            print("=" * 80)
            
        else:
            print("✗ 转换失败")
            print()
            print("可能的原因：")
            print("  1. PDF 文件已损坏")
            print("  2. Ghostscript 未正确安装")
            print("  3. 没有权限读取 PDF 文件")
            print("  4. PDF 文件路径错误")
            print()
            sys.exit(1)
