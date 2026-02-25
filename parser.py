"""
æ‹›æ ‡æ–‡ä»¶è§£ææ¨¡å— - å¢å¼ºç‰ˆ
"""

import re
import json
import os
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import PyPDF2
import docx
import subprocess


class ParseResult:
    """è§£æç»“æœç±»"""
    def __init__(self, requirements, confidence_score=0.0):
        self.requirements = requirements
        self.confidence_score = confidence_score  # 0.0-1.0
        
    def get_confidence_level(self):
        """è·å–ç½®ä¿¡åº¦ç­‰çº§"""
        if self.confidence_score >= 0.8:
            return "é«˜"
        elif self.confidence_score >= 0.6:
            return "ä¸­"
        elif self.confidence_score >= 0.4:
            return "ä½"
        else:
            return "ä¸ç¡®å®š"
    
    def get_confidence_color(self):
        """è·å–ç½®ä¿¡åº¦é¢œè‰²æ ‡è¯†"""
        if self.confidence_score >= 0.8:
            return "ğŸŸ¢"  # é«˜ - ç»¿è‰²
        elif self.confidence_score >= 0.6:
            return "ğŸŸ¡"  # ä¸­ - é»„è‰²
        elif self.confidence_score >= 0.4:
            return "ğŸŸ "  # ä½ - æ©™è‰²
        else:
            return "âšª"  # ä¸ç¡®å®š - ç™½è‰²


class TenderParser:
    """æ‹›æ ‡æ–‡ä»¶è§£æå™¨"""

    def __init__(self, data_dir: Path):
        self.data_dir = data_dir

    def parse_file(self, filepath: Path) -> ParseResult:
        """è§£ææ‹›æ ‡æ–‡ä»¶"""
        
        # æ ¹æ®æ–‡ä»¶ç±»å‹é€‰æ‹©è§£ææ–¹æ³•
        if not filepath.exists():
            return ParseResult([], confidence_score=0.0)

        try:
            if filepath.suffix.lower() == '.pdf':
                return self._parse_pdf(filepath)
            elif filepath.suffix.lower() == '.docx':
                return self._parse_docx(filepath)
            elif filepath.suffix.lower() == '.doc':
                return self._parse_doc(filepath)
            else:
                return ParseResult([], confidence_score=0.0)
        except Exception as e:
            print(f"âœ— æ–‡ä»¶è§£æå¤±è´¥: {e}")
            return ParseResult([], confidence_score=0.0)

    def _parse_pdf(self, filepath: Path) -> ParseResult:
        """è§£æPDFæ–‡ä»¶"""
        print(f"ğŸ“„ å¼€å§‹è§£æPDFæ–‡ä»¶: {filepath.name}")
        
        requirements = []
        
        try:
            reader = PyPDF2.PdfReader(filepath)
            
            # æå–æ–‡æœ¬
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            
            # è§£æéœ€æ±‚
            requirements = self._extract_requirements_from_text(text)
            
            # è®¡ç®—ç½®ä¿¡åº¦
            confidence = self._calculate_confidence(requirements, 'pdf')
            
            print(f"âœ“ PDFè§£æå®Œæˆ")
            print(f"  - æå–éœ€æ±‚: {len(requirements)}")
            print(f"  - ç½®ä¿¡åº¦: {confidence:.2f} ({confidence * 100:.0f}%)")
            
            return ParseResult(requirements, confidence_score=confidence)
            
        except Exception as e:
            print(f"âœ— PDFè§£æå¤±è´¥: {e}")
            return ParseResult([], confidence_score=0.0)

    def _parse_docx(self, filepath: Path) -> ParseResult:
        """è§£æ DOCX æ–‡ä»¶"""
        print(f"ğŸ“„ å¼€å§‹è§£æ DOCX æ–‡ä»¶: {filepath.name}")
        
        requirements = []
        
        try:
            doc = docx.Document(filepath)
            
            # æå–æ–‡æœ¬
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # è§£æéœ€æ±‚
            requirements = self._extract_requirements_from_text(text)
            
            # è®¡ç®—ç½®ä¿¡åº¦
            confidence = self._calculate_confidence(requirements, 'docx')
            
            print(f"âœ“ DOCXè§£æå®Œæˆ")
            print(f"  - æå–éœ€æ±‚: {len(requirements)}")
            print(f"  - ç½®ä¿¡åº¦: {confidence:.2f} ({confidence * 100:.0f}%)")
            
            return ParseResult(requirements, confidence_score=confidence)
            
        except Exception as e:
            print(f"âœ— DOCXè§£æå¤±è´¥: {e}")
            return ParseResult([], confidence_score=0.0)

    def _parse_doc(self, filepath: Path) -> ParseResult:
        """è§£æ DOC æ–‡ä»¶ï¼ˆä½¿ç”¨ antiword è½¬æ¢ï¼‰"""
        print(f"ğŸ“„ å¼€å§‹è§£æ DOC æ–‡ä»¶: {filepath.name}")
        
        try:
            # æ£€æŸ¥ antiword æ˜¯å¦å®‰è£…
            result = subprocess.run(['which', 'antiword'], 
                                 capture_output=True, text=True)
            antiword_path = result.stdout.strip()
            
            if not antiword_path:
                print("âš ï¸  antiword æœªå®‰è£…ï¼Œå°è¯•ä½¿ç”¨å…¶ä»–æ–¹æ³•")
                # å°è¯•ä½¿ç”¨ LibreOffice è½¬æ¢
                libreoffice_result = subprocess.run(['which', 'libreoffice', 'soffice'], 
                                                 capture_output=True, text=True)
                libreoffice_path = libreoffice_result.stdout.strip() or libreoffice_result.stderr.strip()
                
                if libreoffice_path:
                    print(f"âœ“ æ‰¾åˆ° LibreOffice: {libreoffice_path}")
                    # ä½¿ç”¨ LibreOffice è½¬æ¢ DOC ä¸º DOCX
                    temp_dir = Path("temp")
                    temp_dir.mkdir(exist_ok=True)
                    
                    convert_result = subprocess.run(
                        [libreoffice_path, '--headless', '--convert-to', 'docx', 
                         '--outdir', str(temp_dir), str(filepath)],
                        capture_output=True, text=True, timeout=60
                    )
                    
                    if convert_result.returncode == 0:
                        # æŸ¥æ‰¾è½¬æ¢åçš„æ–‡ä»¶
                        temp_docx = temp_dir / filepath.with_suffix('.docx').name
                        if temp_docx.exists():
                            print(f"âœ“ LibreOffice è½¬æ¢æˆåŠŸ")
                            doc = docx.Document(str(temp_docx))
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            requirements = self._extract_requirements_from_text(text)
                            confidence = self._calculate_confidence(requirements, 'doc')
                            
                            # åˆ é™¤ä¸´æ—¶æ–‡ä»¶
                            temp_docx.unlink()
                            
                            print(f"âœ“ DOCè§£æå®Œæˆ")
                            print(f"  - æå–éœ€æ±‚: {len(requirements)}")
                            print(f"  - ç½®ä¿¡åº¦: {confidence:.2f} ({confidence * 100:.0f}%)")
                            
                            return ParseResult(requirements, confidence_score=confidence)
                        else:
                            print("âš ï¸  è½¬æ¢åçš„æ–‡ä»¶æœªæ‰¾åˆ°")
                    else:
                        print(f"âš ï¸  LibreOffice è½¬æ¢å¤±è´¥: {convert_result.stderr}")
                
                # LibreOffice å¤±è´¥ï¼Œè¿”å›ä½ç½®ä¿¡åº¦
                print("âš ï¸  æ— æ³•è§£æ DOC æ–‡ä»¶")
                return ParseResult([], confidence_score=0.0)
            else:
                # ä½¿ç”¨ antiword ç›´æ¥æå–æ–‡æœ¬
                result = subprocess.run([antiword_path, '-t', str(filepath)],
                                     capture_output=True, text=True, timeout=30)
                
                if result.returncode != 0:
                    print(f"âš ï¸  antiword æå–å¤±è´¥: {result.stderr}")
                    return ParseResult([], confidence_score=0.0)
                
                text = result.stdout
                requirements = self._extract_requirements_from_text(text)
                confidence = self._calculate_confidence(requirements, 'doc')
                
                print(f"âœ“ DOCè§£æå®Œæˆ")
                print(f"  - æå–éœ€æ±‚: {len(requirements)}")
                print(f"  - ç½®ä¿¡åº¦: {confidence:.2f} ({confidence * 100:.0f}%)")
                
                return ParseResult(requirements, confidence_score=confidence)
            
        except Exception as e:
            print(f"âœ— DOCè§£æå¤±è´¥: {e}")
            import traceback
            traceback.print_exc()
            return ParseResult([], confidence_score=0.0)

    def _extract_requirements_from_text(self, text: str) -> List[str]:
        """ä»æ–‡æœ¬ä¸­æå–éœ€æ±‚"""
        requirements = []
        seen = set()  # ç”¨äºå»é‡
        
        # å®šä¹‰éœ€æ±‚å…³é”®è¯
        requirement_keywords = [
            "è¦æ±‚", "è§„å®š", "å¿…é¡»", "åº”", "éœ€", "ä¸å¾—",
            "èµ„è´¨", "è¯ä¹¦", "è®¤è¯", "è®¸å¯", "æ‰§ç…§",
            "æ ‡å‡†", "ç¬¦åˆ", "æ»¡è¶³", "è¾¾åˆ°",
            "æŠ€æœ¯", "è®¾å¤‡", "äº§å“", "ææ–™",
            "æ¡ˆä¾‹", "ä¸šç»©", "ç»éªŒ", "å¹´é™", "å¹´",
            "äººå‘˜", "å·¥ç¨‹å¸ˆ", "é¡¹ç›®ç»ç†", "æŠ€æœ¯è´Ÿè´£äºº",
            "ä¿ä¿®", "è´¨ä¿", "æœåŠ¡", "å”®å",
            "é‡‘é¢", "ä»·æ ¼", "æŠ¥ä»·", "è´¹ç”¨",
            "è´¨é‡", "å®‰å…¨", "ç¯ä¿", "ç¯å¢ƒ",
            "å·¥æœŸ", "æ—¶é—´", "äº¤ä»˜", "å®Œå·¥",
            "æ–‡ä»¶", "æŠ¥å‘Š", "æ£€æµ‹", "æµ‹è¯•",
            "å›¾çº¸", "è®¾è®¡", "æ–¹æ¡ˆ",
            "éªŒæ”¶", "è§„èŒƒ", "æ¡ä»¶",
            "è¯ä¹¦ç¼–å·", "è¯ä¹¦ç­‰çº§", "æœ‰æ•ˆæœŸ",
            "æ³¨å†Œèµ„é‡‘", "æ³¨å†Œèµ„æœ¬", "è¥ä¸šé¢",
            "ISO", "9001", "CCC", "CE"
        ]
        
        # å®šä¹‰éœ€è¦è¿‡æ»¤çš„æ— æ„ä¹‰å†…å®¹æ¨¡å¼
        filter_patterns = [
            r'^[ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹å]+[ã€\.]',  # åºå·
            r'^\d+[ã€\.]',  # æ•°å­—åºå·
            r'^\d+\.\d+\.\d+\.\d+',  # IPåœ°å€
            r'^\d{4}-\d{2}-\d{2}',  # æ—¥æœŸ
            r'^\d{11}$',  # ç”µè¯å·ç 
            r'^\w+@\w+\.\w+$',  # é‚®ç®±
            r'^http',  # URL
            r'^www',  # URL
            r'^æµ·è¶Š',  # å…¬å¸åç§°å¼€å¤´
            r'^æ¹–åŒ—',  # åœ°åŒºåç§°å¼€å¤´
            r'^ç”µæ°”',  # å…¬å¸åç§°å¼€å¤´
            r'^å…¬å¸',  # å…¬å¸åç§°
            r'^æ‹›æ ‡æ–‡ä»¶',  # æ–‡æ¡£ç±»å‹
            r'^æŠ•æ ‡æ–‡ä»¶',  # æ–‡æ¡£ç±»å‹
            r'^æŠ€æœ¯æ–‡ä»¶',  # æ–‡æ¡£ç±»å‹
        ]
        
        # åˆ†å‰²æ–‡æœ¬ä¸ºæ®µè½
        paragraphs = text.split('\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            
            # è·³è¿‡ç©ºè¡Œæˆ–å¤ªçŸ­çš„è¡Œ
            if not paragraph or len(paragraph) < 5:
                continue
            
            # è·³è¿‡æ˜æ˜¾æ— æ„ä¹‰çš„è¡Œ
            if any(re.match(pattern, paragraph) for pattern in filter_patterns):
                continue
            
            # æå–åŒ…å«å…³é”®è¯çš„å¥å­
            sentences = re.split(r'[ã€‚ï¼ï¼Ÿï¼›\n]', paragraph)
            
            for sentence in sentences:
                sentence = sentence.strip()
                
                # è·³è¿‡ç©ºè¡Œæˆ–å¤ªçŸ­çš„å¥å­
                if not sentence or len(sentence) < 5:
                    continue
                
                # å»é™¤å¤šä½™ç©ºç™½
                sentence = ' '.join(sentence.split())
                
                # æ£€æŸ¥æ˜¯å¦åŒ…å«éœ€æ±‚å…³é”®è¯
                has_requirement = any(keyword in sentence for keyword in requirement_keywords)
                
                if has_requirement:
                    # å»é‡ï¼šåªä¿ç•™ç¬¬ä¸€æ¬¡å‡ºç°çš„
                    sentence_lower = sentence.lower()
                    if sentence_lower not in seen:
                        seen.add(sentence_lower)
                        requirements.append(sentence)
                        if len(requirements) >= 20:  # æœ€å¤šæå–20ä¸ªéœ€æ±‚
                            break
            
            # å¦‚æœå·²ç»æå–äº†è¶³å¤Ÿçš„éœ€æ±‚ï¼Œåœæ­¢
            if len(requirements) >= 20:
                break
        
        return requirements

    def _calculate_confidence(self, requirements: List[str], file_type: str) -> float:
        """è®¡ç®—è§£æç½®ä¿¡åº¦"""
        if not requirements:
            return 0.0
        
        confidence = 0.0
        
        # 1. åŸºäºéœ€æ±‚æ•°é‡çš„ç½®ä¿¡åº¦ï¼ˆæœ€å¤š 30 åˆ†ï¼‰
        count_score = min(len(requirements) * 1.5, 30)
        
        # 2. åŸºäºéœ€æ±‚è´¨é‡çš„ç½®ä¿¡åº¦ï¼ˆæœ€å¤š 40 åˆ†ï¼‰
        quality_score = 0.0
        for req in requirements[:10]:  # åªæ£€æŸ¥å‰10ä¸ª
            req_lower = req.lower()
            
            # éœ€æ±‚é•¿åº¦
            if len(req) > 10:
                quality_score += 3
            elif len(req) > 20:
                quality_score += 4
            
            # éœ€æ±‚å…·ä½“æ€§
            concrete_keywords = ["èµ„è´¨", "è¯ä¹¦", "ç»éªŒ", "å¹´é™", "å¹´", "çº§", "ISO", "9001", "CCC", "CE"]
            if any(keyword in req_lower for keyword in concrete_keywords):
                quality_score += 2
            
            # éœ€æ±‚æ˜ç¡®æ€§
            vague_keywords = ["ç­‰", "ç›¸å…³", "ç±»ä¼¼", "æœ€å¥½", "åº”", "éœ€"]
            if not any(keyword in req_lower for keyword in vague_keywords):
                quality_score += 2
        
        # å½’ä¸€åŒ–
        quality_score = min(quality_score, 40) / len(requirements[:10]) * 10 if requirements else 0
        
        # 3. åŸºäºæ–‡ä»¶ç±»å‹çš„ç½®ä¿¡åº¦ï¼ˆæœ€å¤š 30 åˆ†ï¼‰
        type_score = {
            'pdf': 30,      # PDF è§£æé€šå¸¸æœ€å¯é 
            'docx': 25,     # DOCX è§£æè¾ƒå¯é 
            'doc': 15       # DOC éœ€è¦è½¬æ¢ï¼Œç½®ä¿¡åº¦è¾ƒä½
        }.get(file_type, 10)
        
        # æ€»ç½®ä¿¡åº¦
        confidence = (count_score + quality_score + type_score) / 100
        
        # ç¡®ä¿ç½®ä¿¡åº¦åœ¨ 0.0-1.0 ä¹‹é—´
        return max(0.0, min(confidence, 1.0))
    
    def _get_suggestions(self, parse_result: ParseResult) -> List[str]:
        """ç”Ÿæˆæ”¹è¿›å»ºè®®"""
        suggestions = []
        level = parse_result.get_confidence_level()
        
        if level == "ä½" or level == "ä¸ç¡®å®š":
            suggestions.append("æ–‡ä»¶å¯èƒ½ä¸æ˜¯æ ‡å‡†æ‹›æ ‡æ–‡ä»¶æ ¼å¼ï¼Œè¯·æ£€æŸ¥æ–‡ä»¶å†…å®¹")
            suggestions.append("å»ºè®®å°†æ–‡ä»¶è½¬æ¢ä¸º PDF æˆ– DOCX æ ¼å¼åé‡æ–°ä¸Šä¼ ")
            suggestions.append("å¯ä»¥å°è¯•æ‰‹åŠ¨è¾“å…¥éœ€æ±‚")
        
        if len(parse_result.requirements) < 5:
            suggestions.append("æå–çš„éœ€æ±‚è¾ƒå°‘ï¼Œå¯èƒ½é—æ¼äº†éƒ¨åˆ†å†…å®¹")
            suggestions.append("å»ºè®®äººå·¥è¡¥å……é‡è¦çš„éœ€æ±‚")
        
        # æ£€æŸ¥éœ€æ±‚è´¨é‡
        vague_count = 0
        for req in parse_result.requirements[:10]:
            if any(kw in req.lower() for kw in ["ç­‰", "ç›¸å…³", "ç±»ä¼¼", "æœ€å¥½"]):
                vague_count += 1
        
        if vague_count > 2:
            suggestions.append("éƒ¨åˆ†éœ€æ±‚è¡¨è¾¾ä¸å¤Ÿå…·ä½“ï¼Œå»ºè®®æ˜ç¡®åŒ–")
        
        return suggestions


# æµ‹è¯•ä»£ç 
if __name__ == "__main__":
    import sys
    import time
    
    # æµ‹è¯•è§£æ
    test_dir = Path(__file__).parent / "tests"
    test_dir.mkdir(exist_ok=True)
    
    # åˆ›å»ºæµ‹è¯•æ–‡ä»¶
    print("åˆ›å»ºæµ‹è¯•æ–‡ä»¶...")
    test_pdf = test_dir / "test_tender.pdf"
    test_docx = test_dir / "test_tender.docx"
    test_doc = test_dir / "test_tender.doc"
    
    # åˆ›å»ºç®€å•çš„æµ‹è¯•æ–‡ä»¶ï¼ˆå¦‚æœä¸å­˜åœ¨ï¼‰
    if not test_docx.exists():
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("æµ‹è¯•æ‹›æ ‡æ–‡ä»¶", level=1)
            doc.add_paragraph("é¡¹ç›®åç§°ï¼šæŸå·¥ä¸šå›­åŒº10kVå¼€å…³æŸœé‡‡è´­")
            doc.add_paragraph("å®¢æˆ·ï¼šæŸå·¥ä¸šå›­åŒº")
            doc.add_paragraph("æˆªæ­¢æ—¥æœŸï¼š2026-03-15")
            doc.add_paragraph("")
            doc.add_heading("æŠ€æœ¯è¦æ±‚", level=2)
            doc.add_paragraph("1. äº§å“è¦æ±‚")
            doc.add_paragraph("   - KYN28A-12 æˆ·å†…äº¤æµé‡‘å±é“ è£…ç§»å¼€å¼å¼€å…³è®¾å¤‡", style="List Bullet")
            doc.add_paragraph("   - é¢å®šç”µå‹ï¼š10kV", style="List Bullet")
            doc.add_paragraph("   - é¢å®šç”µæµï¼š630A", style="List Bullet")
            doc.add_paragraph("   - é˜²æŠ¤ç­‰çº§ï¼šIP30", style="List Bullet")
            doc.add_paragraph("")
            doc.add_paragraph("2. èµ„è´¨è¦æ±‚")
            doc.add_paragraph("   - ç”µåŠ›å·¥ç¨‹æ–½å·¥æ€»æ‰¿åŒ…ä¸‰çº§åŠä»¥ä¸Š", style="List Bullet")
            doc.add_paragraph("   - è´¨é‡ç®¡ç†ä½“ç³»è®¤è¯", style="List Bullet")
            doc.add_paragraph("   - ç¯å¢ƒç®¡ç†ä½“ç³»è®¤è¯", style="List Bullet")
            doc.add_paragraph("")
            doc.add_paragraph("3. ä¸šç»©è¦æ±‚")
            doc.add_paragraph("   - æä¾›ç±»ä¼¼é¡¹ç›®æ¡ˆä¾‹3ä¸ª", style="List Bullet")
            doc.add_paragraph("   - é¡¹ç›®é‡‘é¢åœ¨50ä¸‡å…ƒä»¥ä¸Š", style="List Bullet")
            doc.add_paragraph("   - é¡¹ç›®ç»éªŒ5å¹´ä»¥ä¸Š", style="List Bullet")
            doc.add_paragraph("")
            doc.add_heading("å•†åŠ¡è¦æ±‚", level=2)
            doc.add_paragraph("1. æŠ¥ä»·è¦æ±‚")
            doc.add_paragraph("   - å›ºå®šæ€»ä»·", style="List Bullet")
            doc.add_paragraph("   - æŠ¥ä»·æœ‰æ•ˆæœŸ30å¤©", style="List Bullet")
            doc.add_paragraph("")
            doc.add_paragraph("2. ä»˜æ¬¾è¦æ±‚")
            doc.add_paragraph("   - éªŒæ”¶å90å¤©ä»˜æ¬¾", style="List Bullet")
            
            doc.save(str(test_docx))
            print(f"âœ“ åˆ›å»ºæµ‹è¯• DOCX æ–‡ä»¶: {test_docx}")
        except Exception as e:
            print(f"âœ— åˆ›å»ºæµ‹è¯•æ–‡ä»¶å¤±è´¥: {e}")
    
    # æµ‹è¯•è§£æ
    parser = TenderParser(Path(__file__).parent / "data")
    
    print("\n" + "=" * 60)
    print("å¼€å§‹æµ‹è¯•è§£æå™¨")
    print("=" * 60 + "\n")
    
    if test_docx.exists():
        print("\næµ‹è¯• 1ï¼šè§£æ DOCX æ–‡ä»¶")
        result = parser.parse_file(test_docx)
        print(f"ç½®ä¿¡åº¦ï¼š{result.confidence_score:.2f}")
        print(f"ç­‰çº§ï¼š{result.get_confidence_level()}")
    
    time.sleep(1)
    
    print("\nâœ“ æ‰€æœ‰æµ‹è¯•å®Œæˆ")
