

    def _add_qualifications_with_images(self, doc: Document, qualifications: List[Dict], data_dir: Path):
        """
        添加企业资质（PDF转图片）

        优点：
        - 打印出来直接可以看到证书图片
        - 客户不需要任何操作
        - 投标文件看起来更专业

        缺点：
        - 需要安装额外库（pdf2image + Pillow + Ghostscript）
        - 文件会大一些
        """
        if not PDF_TO_IMAGE_AVAILABLE or PILImage is None:
            # PDF转图片功能不可用，使用原方法
            self._add_qualifications_v2(doc, qualifications)
            return

        p = doc.add_paragraph()
        run = p.add_run("第3章 企业资质")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        if not qualifications:
            doc.add_paragraph("（具体资质文件详见附件）")
            doc.add_page_break()
            return

        # 3.1 体系认证证书
        p = doc.add_paragraph()
        run = p.add_run("3.1 体系认证证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出体系认证
        system_certs = [q for q in qualifications if "体系" in q.get("name", "") or "认证" in q.get("name", "")]

        for cert in system_certs[:10]:  # 最多显示10个
            if cert.get('cert_file'):
                cert_path = data_dir / cert['cert_file']
                
                # 显示证书名称
                p = doc.add_paragraph()
                run = p.add_run(f"• {cert['name']}（{cert['level']}）")
                run.bold = True
                
                # 显示证书编号和有效期
                if cert.get('cert_no'):
                    p.add_paragraph(f"  证书编号：{cert['cert_no']}")
                if cert.get('valid_until'):
                    p.add_paragraph(f"  有效期至：{cert['valid_until']}")
                
                # PDF转图片并插入
                if cert_path.exists():
                    try:
                        import tempfile
                        with tempfile.TemporaryDirectory() as temp_dir:
                            # 转换PDF为图片（只转换第一页）
                            images = convert_from_path(
                                str(cert_path),
                                output_folder=temp_dir,
                                first_page_only=True,
                                dpi=200,  # 分辨率，200dpi打印效果较好
                                fmt='jpg'  # 输出为JPG格式
                                use_cropbox=True
                            )
                            
                            if images:
                                # 调整图片大小
                                img = PILImage.open(images[0])
                                img_width, img_height = img.size
                                
                                # 如果图片太宽，调整宽度
                                target_width = 500  # 目标宽度500像素
                                if img_width > target_width:
                                    ratio = target_width / img_width
                                    new_height = int(img_height * ratio)
                                    img = img.resize((target_width, new_height), PILImage.LANCZOS)
                                    
                                    # 保存调整后的图片
                                    img.save(images[0], quality=85)
                                
                                # 插入图片到Word文档
                                doc.add_paragraph()  # 空行
                                doc.add_picture(images[0], width=Inches(5.5))  # 宽度5.5英寸
                                doc.add_paragraph()  # 空行
                    except Exception as e:
                        # 如果转换失败，显示文件路径
                        p.add_paragraph(f"  证书文件详见附件：{cert['cert_file']}")
                        p.add_paragraph(f"  （图片转换失败：{str(e)}）")
                else:
                    p.add_paragraph(f"  证书文件：{cert['cert_file']}（文件不存在）")
        
        doc.add_paragraph()  # 证书之间空一行
"""
投标文件生成模块 - V2（短期升级版）
"""

from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# PDF 转图片相关导入（需要安装 pdf2image 和 Pillow）
try:
    from pdf2image import convert_from_path
    from PIL import Image as PILImage
    PDF_TO_IMAGE_AVAILABLE = True
except ImportError:
    PDF_TO_IMAGE_AVAILABLE = False
    print("警告: pdf2image 或 Pillow 未安装，PDF转图片功能将不可用")
    print("安装命令: pip install pdf2image Pillow")
    # 添加导入检查（避免导入错误）
    try:
        from PIL import Image as PILImage
    except ImportError:
        PILImage = None
        PDF_TO_IMAGE_AVAILABLE = False


class BidDocumentGenerator:
    """投标文件生成器 - V2"""

    def __init__(self, templates_dir: Path, output_dir: Path):
        self.templates_dir = templates_dir
        self.output_dir = output_dir

    def generate_bid(self, tender_info: Dict, company_info: Dict,
                    matched_data: Dict, quote_data: Dict = None) -> Path:
        """
        生成投标文件

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据（资质、案例、产品等）
            quote_data: 报价数据

        Returns:
            生成的文件路径
        """
        # 创建新文档
        doc = Document()

        # 生成各个章节
        self._add_cover_v2(doc, tender_info, company_info, bid_type="投标文件")
        self._add_company_proof(doc, company_info)
        self._add_bid纲领_v2(doc, company_info, tender_info)
        self._add_deviation_table(doc, tender_info, table_type="技术")  # 新增
        self._add_company_intro_v2(doc, company_info)  # 升级
        self._add_tech_solution(doc, tender_info, matched_data)
        self._add_quotation(doc, quote_data if quote_data else {})
        self._add_qualifications_with_images(doc, matched_data.get("qualifications", []), DATA_DIR)  # 升级
        self._add_performance(doc, matched_data.get("cases", []))
        self._add_after_sales(doc, company_info)
        self._add_response_commitment(doc)  # 新增
        self._add_commitment(doc)

        # 保存文件
        project_name = tender_info.get("project_info", {}).get("project_name", "未知项目")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"投标文件_{project_name}_{timestamp}.docx"

        # 清理文件名
        filename = "".join(c for c in filename if c not in '\/:*?"<>|')

        output_path = self.output_dir / filename
        doc.save(output_path)

        return output_path

    def generate_separate_bids(self, tender_info: Dict, company_info: Dict,
                               matched_data: Dict, quote_data: Dict = None) -> Dict[str, Path]:
        """
        生成分开的技术标和商务标

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据（资质、案例、产品等）
            quote_data: 报价数据

        Returns:
            包含技术标和商务标路径的字典
        """
        # 生成技术标
        tech_path = self.generate_tech_bid(tender_info, company_info, matched_data)

        # 生成商务标
        commercial_path = self.generate_commercial_bid(
            tender_info, company_info, matched_data, quote_data
        )

        return {
            "tech": tech_path,
            "commercial": commercial_path
        }

    def generate_tech_bid(self, tender_info: Dict, company_info: Dict,
                          matched_data: Dict) -> Path:
        """
        生成技术标

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据

        Returns:
            生成的文件路径
        """
        # 创建新文档
        doc = Document()

        # 生成技术标章节
        self._add_cover_v2(doc, tender_info, company_info, bid_type="技术投标文件")
        self._add_company_proof(doc, company_info)
        self._add_bid纲领_v2(doc, company_info, tender_info)
        self._add_deviation_table(doc, tender_info, table_type="技术")  # 新增
        self._add_company_intro_v2(doc, company_info)  # 升级
        self._add_tech_solution(doc, tender_info, matched_data)
        self._add_qualifications_with_images(doc, matched_data.get("qualifications", []), DATA_DIR)  # 升级
        self._add_performance(doc, matched_data.get("cases", []))
        self._add_tech_commitment(doc)
        self._add_response_commitment(doc)  # 新增

        # 保存文件
        project_name = tender_info.get("project_info", {}).get("project_name", "未知项目")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"技术标_{project_name}_{timestamp}.docx"

        # 清理文件名
        filename = "".join(c for c in filename if c not in '\/:*?"<>|')

        output_path = self.output_dir / filename
        doc.save(output_path)

        return output_path

    def generate_commercial_bid(self, tender_info: Dict, company_info: Dict,
                                matched_data: Dict, quote_data: Dict = None) -> Path:
        """
        生成商务标

        Args:
            tender_info: 招标信息
            company_info: 公司信息
            matched_data: 匹配的数据
            quote_data: 报价数据

        Returns:
            生成的文件路径
        """
        # 创建新文档
        doc = Document()

        # 生成商务标章节
        self._add_cover_v2(doc, tender_info, company_info, bid_type="商务投标文件")
        self._add_company_proof(doc, company_info)
        self._add_bid纲领_v2(doc, company_info, tender_info)
        self._add_deviation_table(doc, tender_info, table_type="商务")  # 新增
        self._add_company_intro_v2(doc, company_info)  # 升级
        self._add_response_commitment(doc)  # 新增
        self._add_quotation(doc, quote_data if quote_data else {})
        self._add_qualifications_with_images(doc, matched_data.get("qualifications", []), DATA_DIR)  # 升级
        self._add_performance(doc, matched_data.get("cases", []))
        self._add_after_sales(doc, company_info)
        self._add_commercial_commitment(doc)

        # 保存文件
        project_name = tender_info.get("project_info", {}).get("project_name", "未知项目")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"商务标_{project_name}_{timestamp}.docx"

        # 清理文件名
        filename = "".join(c for c in filename if c not in '\/:*?"<>|')

        output_path = self.output_dir / filename
        doc.save(output_path)

        return output_path

    def _add_cover_v2(self, doc: Document, tender_info: Dict, company_info: Dict, bid_type: str = "投标文件"):
        """
        添加封面 - V2（符合公司标准格式）

        格式：
        项目编号：xxx
        [招标人]
        [项目名称]
        [投标文件类型]

        申请人：海越(湖北)电气股份有限公司（盖单位章）
        [年]年[月]月[日]日
        """
        project_info = tender_info.get("project_info", {})
        project_name = project_info.get("project_name", "")
        project_no = project_info.get("project_no", "未知项目编号")
        tenderer = project_info.get("tenderer", "未知招标人")

        # 公司名称和英文名
        for _ in range(3):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company_info["name"])
        run.bold = True
        run.font.size = Pt(24)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company_info["name_en"])
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        # 地址和联系方式
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"地址：{company_info['address']}  P.C.:{company_info['postal_code']}")
        run.font.size = Pt(9)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"E-mail:{company_info['email']} Tel:{company_info['phone']} Fax:{company_info['fax']}")
        run.font.size = Pt(9)

        # 分隔线
        doc.add_paragraph("-" * 50)

        # 项目编号（标准格式）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"项目编号：{project_no}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 招标人
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(tenderer)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 项目名称
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_name)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = "宋体"

        doc.add_paragraph()
        doc.add_paragraph()

        # 投标文件类型
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(bid_type)
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "黑体"

        # 日期和签字
        for _ in range(4):
            doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run(f"申请人： {company_info['name']} （盖单位章）")
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        run = p.add_run("法定代表人或其委托代理人： （签字）")
        run.font.size = Pt(12)

        today = datetime.now()
        p = doc.add_paragraph()
        run = p.add_run(f"{today.year} 年 {today.month} 月 {today.day} 日")
        run.font.size = Pt(12)

        doc.add_page_break()

    def _add_company_proof(self, doc: Document, company_info: Dict):
        """添加公司关系证明"""
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("海越(湖北)电气股份有限公司 公司关系证明")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph(f"地址：{company_info['address']}  P.C.:{company_info['postal_code']}")
        doc.add_paragraph(f"E-mail:{company_info['email']} Tel:{company_info['phone']} Fax:{company_info['fax']}")

        doc.add_paragraph("-" * 50)

        doc.add_page_break()

    def _add_bid纲领_v2(self, doc: Document, company_info: Dict, tender_info: Dict):
        """
        添加投标纲领 - V2（符合公司标准格式）

        包含：
        - 标题
        - 致辞
        - 公司介绍（两大生产基地）
        - 行业经验和合作伙伴
        - 承诺
        """
        project_info = tender_info.get("project_info", {})

        # 标题
        p = doc.add_paragraph()
        run = p.add_run("项目投标纲领")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph(f"地址：{company_info['address']}  P.C.:{company_info['postal_code']}")
        doc.add_paragraph(f"E-mail:{company_info['email']} Tel:{company_info['phone']} Fax:{company_info['fax']}")

        doc.add_paragraph("-" * 50)

        doc.add_paragraph("项目 投 标 纲 领")

        p = doc.add_paragraph()
        run = p.add_run("尊敬的各位领导、专家评委：你们好！")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        tenderer = project_info.get("tenderer", "贵公司")
        project_name = project_info.get("project_name", "本项目")

        doc.add_paragraph(f"非常荣幸得以参与{tenderer}{project_name}的投标。")
        doc.add_paragraph()

        doc.add_paragraph(
            "本公司在宁波及湖北各设有一处生产基地，其中湖北生产基地位于湖北省孝感市孝昌县"
            "经济开发区华阳大道 188-1 号，宁波生产基地则位于宁波市镇海骆驼机电园区盛兴路 348 号。"
            "本公司是一家专注于制造各类配电设备的实体型企业，具备稳定的客户基础及卓越的"
            "市场拓展能力，年产值可达 20 亿元人民币，年上缴税收 1 亿元人民币，并计划在 5 年内完成上市。"
        )

        doc.add_paragraph()
        doc.add_paragraph(
            "本公司在高低压成套配电设备行业深耕多年，无论是海越品牌高低压成套配电设备，还是"
            "三大品牌（ABB、西门子、施耐德）授权柜，本公司都有着丰富的制造经验，并且与许多知名企业"
            "都有着深入的合作。"
        )

        doc.add_paragraph()
        doc.add_paragraph("我司将以优质的产品、合理的价格、完善的服务，竭诚为广大客户提供最满意的配电解决方案！")

        doc.add_page_break()

    def _add_deviation_table(self, doc: Document, tender_info: Dict, table_type: str = "技术"):
        """
        添加偏离表 - 新增

        Args:
            doc: Word文档对象
            tender_info: 招标信息
            table_type: 偏离表类型（"技术"或"商务"）
        """
        # 标题
        p = doc.add_paragraph()
        run = p.add_run(f"第1章 {table_type}偏离表")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 说明
        doc.add_paragraph("说明：")
        doc.add_paragraph(f"本{table_type}投标文件完全满足招标文件中{table_type}规范书及{table_type}条款的全部要求，无偏离。")
        doc.add_paragraph()

        # 创建表格
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ["序号", "条款编号", "偏离说明"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # 设置表头格式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(12)
            # 【修复】设置单元格背景色
            # 使用 docx 的 background_color 属性（简化版）
            from docx.shared import RGBColor
            cell.background_color = RGBColor(217, 226, 243)  # 浅灰色背景

        # 第一行示例
        cells = table.add_row().cells
        cells[0].text = "1"
        cells[1].text = "-"
        cells[2].text = f"无{table_type}偏离"

        doc.add_paragraph()
        doc.add_paragraph(f"我方承诺：{table_type}投标文件中所投产品完全满足技术规范书中星号条款（*）中的相关要求。")

        doc.add_page_break()

    def _add_company_intro_v2(self, doc: Document, company_info: Dict):
        """
        添加公司介绍 - V2（升级版）

        包含：
        - 企业简介
        - 发展历程
        """
        # 标题
        p = doc.add_paragraph()
        run = p.add_run("企业简介")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 公司介绍
        doc.add_paragraph("海越（湖北）电气股份有限公司是一家专注于高低压成套配电设备研发、制造、销售和服务的高新技术企业。")
        doc.add_paragraph()

        doc.add_paragraph("公司拥有两大生产基地，分别位于湖北孝感和浙江宁波，占地面积约50000平方米，建筑面积约35000平方米。")
        doc.add_paragraph()

        doc.add_paragraph(
            "公司主要产品包括：35kV及以下高压开关柜、低压开关柜、箱式变电站、预制舱式变电站、"
            "配电箱、母线桥等配电设备，广泛应用于电力、钢铁、化工、交通、建筑、市政等行业。"
        )
        doc.add_paragraph()

        doc.add_paragraph("公司通过了ISO9001质量管理体系认证、ISO14001环境管理体系认证、OHSAS18001职业健康安全管理体系认证。")
        doc.add_paragraph()

        doc.add_paragraph("公司拥有完善的研发、生产、销售和服务体系，具备从方案设计、产品制造到安装调试、售后服务的全流程能力。")
        doc.add_paragraph()

        # 发展历程
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("发展历程")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        milestones = [
            "• 成立初期：公司成立，专注于高低压配电设备研发制造",
            "• 业务扩展：建立湖北孝感生产基地，扩大生产规模",
            "• 市场拓展：建立浙江宁波生产基地，形成双基地格局",
            "• 技术升级：引进先进生产设备，提升产品质量",
            "• 品牌合作：与ABB、西门子等国际知名品牌建立合作关系",
            "• 规模扩张：年产值突破20亿元，成为行业领先企业",
            "• 未来规划：持续技术创新，计划5年内完成上市"
        ]

        for milestone in milestones:
            doc.add_paragraph(milestone)

        doc.add_page_break()

    def _add_tech_solution(self, doc: Document, tender_info: Dict, matched_data: Dict):
        """添加技术方案"""
        p = doc.add_paragraph()
        run = p.add_run("第2章 方案建议书")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 2.1 工艺质量
        p = doc.add_paragraph()
        run = p.add_run("2.1 工艺质量")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        product_reqs = tender_info.get("product_requirements", [])

        doc.add_paragraph("根据招标文件要求，我司为贵方提供一套完整、可靠、先进的配电设备解决方案。")
        doc.add_paragraph()

        if product_reqs:
            doc.add_paragraph("主要产品配置：")
            for i, product in enumerate(product_reqs[:10], 1):  # 最多显示10个
                doc.add_paragraph(f"{i}. {product}")
            doc.add_paragraph()

        # 2.2 技术特点
        p = doc.add_paragraph()
        run = p.add_run("2.2 技术特点")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        features = [
            "（1）产品符合国家相关标准和规范",
            "（2）采用优质元器件，确保设备可靠运行",
            "（3）结构设计合理，便于安装和维护",
            "（4）防护等级高，适应各种环境条件",
            "（5）完善的五防联锁功能，运行安全可靠",
            "（6）采用真空断路器，性能优越，维护方便"
        ]

        for feature in features:
            doc.add_paragraph(feature)

        doc.add_paragraph()

        # 2.3 主要元器件品牌
        p = doc.add_paragraph()
        run = p.add_run("2.3 主要元器件品牌")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        brands = [
            "• 真空断路器：上海人民、厦门华电",
            "• 保护装置：西门子、ABB、施耐德",
            "• 电流互感器：大连第一互感器",
            "• 断路器：施耐德、ABB",
            "• 仪表：斯菲尔、许继",
            "• 接触器：施耐德、ABB",
            "• 继电器：ABB、西门子"
        ]

        for brand in brands:
            doc.add_paragraph(brand)

        doc.add_page_break()

    def _add_quotation(self, doc: Document, quote_data: Dict):
        """添加报价单"""
        p = doc.add_paragraph()
        run = p.add_run("报价明细")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        # 报价汇总
        products = quote_data.get("products", [])
        total_amount = sum(p.get("小计", 0) for p in products)

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ["序号", "产品名称", "型号/规格", "单位", "数量", "单价(元)"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # 填充数据
        for i, product in enumerate(products, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = product.get("产品名称", "")
            row_cells[2].text = product.get("型号/规格", "")
            row_cells[3].text = product.get("单位", "")
            row_cells[4].text = str(product.get("数量", ""))
            row_cells[5].text = str(product.get("单价", ""))

        doc.add_paragraph()

        # 合计
        p = doc.add_paragraph()
        run = p.add_run(f"报价合计：人民币 {total_amount:,.2f} 元")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 0, 0)

        # 报价说明
        doc.add_paragraph()
        doc.add_paragraph("报价说明：")
        doc.add_paragraph("1. 以上报价含13%增值税，含国内运费")
        doc.add_paragraph("2. 交货期：合同签订后30天内（可根据要求调整）")
        doc.add_paragraph("3. 质量保证：产品免费保修一年，终身提供技术服务")
        doc.add_paragraph("4. 付款方式：按合同约定执行")

        doc.add_page_break()

    def _add_qualifications_v2(self, doc: Document, qualifications: List[Dict]):
        """
        添加资质文件 - V2（升级版）

        包含：
        - 资质列表（表格形式）
        - 证书文件路径说明
        - 支持的证明材料
        """
        p = doc.add_paragraph()
        run = p.add_run("第3章 资质证明")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        if not qualifications:
            doc.add_paragraph("（具体资质文件详见附件）")
            doc.add_page_break()
            return

        # 3.1 体系认证证书
        p = doc.add_paragraph()
        run = p.add_run("3.1 体系认证证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出体系认证
        system_certs = [q for q in qualifications if "体系" in q.get("name", "") or "认证" in q.get("name", "")]

        for cert in system_certs[:10]:  # 最多显示10个
            doc.add_paragraph(f"• {cert['name']}（{cert['level']}）")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.2 资信等级证书
        p = doc.add_paragraph()
        run = p.add_run("3.2 资信等级证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出资信等级
        credit_certs = [q for q in qualifications if "AAA" in q.get("name", "") or "资信" in q.get("name", "")]

        for cert in credit_certs[:10]:
            doc.add_paragraph(f"• {cert['name']}")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.3 省级荣誉证书
        p = doc.add_paragraph()
        run = p.add_run("3.3 省级荣誉证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出省级荣誉
        honor_certs = [q for q in qualifications if "重点" in q.get("name", "") or "质量奖" in q.get("name", "")]

        for cert in honor_certs[:10]:
            doc.add_paragraph(f"• {cert['name']}")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.4 合作伙伴证书
        p = doc.add_paragraph()
        run = p.add_run("3.4 合作伙伴证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出合作伙伴
        partner_certs = [q for q in qualifications if "授权" in q.get("name", "") or "合作" in q.get("name", "")]

        for cert in partner_certs[:10]:
            doc.add_paragraph(f"• {cert['name']}")
            if cert.get('cert_file'):
                doc.add_paragraph(f"  证书文件：{cert['cert_file']}")
        doc.add_paragraph()

        # 3.5 其他证书
        remaining = len(qualifications) - len(system_certs) - len(credit_certs) - len(honor_certs) - len(partner_certs)
        if remaining > 0:
            p = doc.add_paragraph()
            run = p.add_run("3.5 其他证书")
            run.bold = True

    def _add_qualifications_with_images(self, doc: Document, qualifications: List[Dict], data_dir: Path):
        """
        添加企业资质（PDF转图片）

        优点：
        - 打印出来直接可以看到证书图片
        - 客户不需要任何操作
        - 投标文件看起来更专业

        缺点：
        - 需要安装额外库（pdf2image + Pillow + Ghostscript）
        - 文件会大一些
        """
        if not PDF_TO_IMAGE_AVAILABLE or PILImage is None:
            # PDF转图片功能不可用，使用原方法
            self._add_qualifications_v2(doc, qualifications)
            return

        p = doc.add_paragraph()
        run = p.add_run("第3章 企业资质")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        if not qualifications:
            doc.add_paragraph("（具体资质文件详见附件）")
            doc.add_page_break()
            return

        # 3.1 体系认证证书
        p = doc.add_paragraph()
        run = p.add_run("3.1 体系认证证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出体系认证
        system_certs = [q for q in qualifications if "体系" in q.get("name", "") or "认证" in q.get("name", "")]

        for cert in system_certs[:10]:  # 最多显示10个
            if cert.get('cert_file'):
                cert_path = data_dir / cert['cert_file']
                
                # 显示证书名称
                p = doc.add_paragraph()
                run = p.add_run(f"• {cert['name']}（{cert['level']}）")
                run.bold = True
                
                # 显示证书编号和有效期
                if cert.get('cert_no'):
                    p.add_paragraph(f"  证书编号：{cert['cert_no']}")
                if cert.get('valid_until'):
                    p.add_paragraph(f"  有效期至：{cert['valid_until']}")
                
                # PDF转图片并插入
                if cert_path.exists():
                    try:
                        import tempfile
                        with tempfile.TemporaryDirectory() as temp_dir:
                            # 转换PDF为图片（只转换第一页）
                            images = convert_from_path(
                                str(cert_path),
                                output_folder=temp_dir,
                                first_page_only=True,
                                dpi=200,  # 分辨率，200dpi打印效果较好
                                fmt='jpg',  # 输出为JPG格式
                                use_cropbox=True
                            )
                            
                            if images:
                                # 调整图片大小
                                img = PILImage.open(images[0])
                                img_width, img_height = img.size
                                
                                # 如果图片太宽，调整宽度
                                target_width = 500  # 目标宽度500像素
                                if img_width > target_width:
                                    ratio = target_width / img_width
                                    new_height = int(img_height * ratio)
                                    img = img.resize((target_width, new_height), PILImage.LANCZOS)
                                    
                                    # 保存调整后的图片
                                    img.save(images[0], quality=85)
                                
                                # 插入图片到Word文档
                                doc.add_paragraph()  # 空行
                                doc.add_picture(images[0], width=Inches(5.5))  # 宽度5.5英寸
                                doc.add_paragraph()  # 空行
                    except Exception as e:
                        # 如果转换失败，显示文件路径
                        p.add_paragraph(f"  证书文件详见附件：{cert['cert_file']}")
                        p.add_paragraph(f"  （图片转换失败：{str(e)}）")
                else:
                    p.add_paragraph(f"  证书文件：{cert['cert_file']}（文件不存在）")
        
        doc.add_paragraph()  # 证书之间空一行

        # 3.2 信用等级证书
        p = doc.add_paragraph()
        run = p.add_run("3.2 信用等级证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出信用等级
        credit_certs = [q for q in qualifications if "AAA" in q.get("name", "") or "信用" in q.get("name", "")]

        for cert in credit_certs[:10]:
            if cert.get('cert_file'):
                cert_path = data_dir / cert['cert_file']
                
                p = doc.add_paragraph()
                run = p.add_run(f"• {cert['name']}")
                run.bold = True
                
                if cert.get('cert_no'):
                    p.add_paragraph(f"  证书编号：{cert['cert_no']}")
                if cert.get('valid_until'):
                    p.add_paragraph(f"  有效期至：{cert['valid_until']}")
                
                # PDF转图片并插入
                if cert_path.exists():
                    try:
                        import tempfile
                        with tempfile.TemporaryDirectory() as temp_dir:
                            images = convert_from_path(
                                str(cert_path),
                                output_folder=temp_dir,
                                first_page_only=True,
                                dpi=200,
                                fmt='jpg',
                                use_cropbox=True
                            )
                            
                            if images:
                                img = PILImage.open(images[0])
                                img_width, img_height = img.size
                                
                                target_width = 500
                                if img_width > target_width:
                                    ratio = target_width / img_width
                                    new_height = int(img_height * ratio)
                                    img = img.resize((target_width, new_height), PILImage.LANCZOS)
                                    img.save(images[0], quality=85)
                                
                                doc.add_paragraph()
                                doc.add_picture(images[0], width=Inches(5.5))
                                doc.add_paragraph()
                    except Exception as e:
                        p.add_paragraph(f"  证书文件详见附件：{cert['cert_file']}")
                        p.add_paragraph(f"  （图片转换失败：{str(e)}）")
        
        doc.add_paragraph()

        # 3.3 省级荣誉证书
        p = doc.add_paragraph()
        run = p.add_run("3.3 省级荣誉证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出省级荣誉
        honor_certs = [q for q in qualifications if "重点" in q.get("name", "") or "质量奖" in q.get("name", "")]

        for cert in honor_certs[:10]:
            if cert.get('cert_file'):
                cert_path = data_dir / cert['cert_file']
                
                p = doc.add_paragraph()
                run = p.add_run(f"• {cert['name']}")
                run.bold = True
                
                if cert.get('cert_no'):
                    p.add_paragraph(f"  证书编号：{cert['cert_no']}")
                if cert.get('valid_until'):
                    p.add_paragraph(f"  有效期至：{cert['valid_until']}")
                
                # PDF转图片并插入
                if cert_path.exists():
                    try:
                        import tempfile
                        with tempfile.TemporaryDirectory() as temp_dir:
                            images = convert_from_path(
                                str(cert_path),
                                output_folder=temp_dir,
                                first_page_only=True,
                                dpi=200,
                                fmt='jpg',
                                use_cropbox=True
                            )
                            
                            if images:
                                img = PILImage.open(images[0])
                                img_width, img_height = img.size
                                
                                target_width = 500
                                if img_width > target_width:
                                    ratio = target_width / img_width
                                    new_height = int(img_height * ratio)
                                    img = img.resize((target_width, new_height), PILImage.LANCZOS)
                                    img.save(images[0], quality=85)
                                
                                doc.add_paragraph()
                                doc.add_picture(images[0], width=Inches(5.5))
                                doc.add_paragraph()
                    except Exception as e:
                        p.add_paragraph(f"  证书文件详见附件：{cert['cert_file']}")
                        p.add_paragraph(f"  （图片转换失败：{str(e)}）")
        
        doc.add_paragraph()

        # 3.4 合作伙伴证书
        p = doc.add_paragraph()
        run = p.add_run("3.4 合作伙伴证书")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"

        doc.add_paragraph()

        # 过滤出合作伙伴
        partner_certs = [q for q in qualifications if "授权" in q.get("name", "") or "合作" in q.get("name", "")]

        for cert in partner_certs[:10]:
            if cert.get('cert_file'):
                cert_path = data_dir / cert['cert_file']
                
                p = doc.add_paragraph()
                run = p.add_run(f"• {cert['name']}")
                run.bold = True
                
                # PDF转图片并插入
                if cert_path.exists():
                    try:
                        import tempfile
                        with tempfile.TemporaryDirectory() as temp_dir:
                            images = convert_from_path(
                                str(cert_path),
                                output_folder=temp_dir,
                                first_page_only=True,
                                dpi=200,
                                fmt='jpg',
                                use_cropbox=True
                            )
                            
                            if images:
                                img = PILImage.open(images[0])
                                img_width, img_height = img.size
                                
                                target_width = 500
                                if img_width > target_width:
                                    ratio = target_width / img_width
                                    new_height = int(img_height * ratio)
                                    img = img.resize((target_width, new_height), PILImage.LANCZOS)
                                    img.save(images[0], quality=85)
                                
                                doc.add_paragraph()
                                doc.add_picture(images[0], width=Inches(5.5))
                                doc.add_paragraph()
                    except Exception as e:
                        p.add_paragraph(f"  证书文件详见附件：{cert['cert_file']}")
                        p.add_paragraph(f"  （图片转换失败：{str(e)}）")
        
        doc.add_paragraph()

        # 3.5 其他证书
        remaining = len(qualifications) - len(system_certs) - len(credit_certs) - len(honor_certs) - len(partner_certs)
        if remaining > 0:
            p = doc.add_paragraph()
            run = p.add_run("3.5 其他证书")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "宋体"

            doc.add_paragraph()
            doc.add_paragraph(f"（其他证书共 {remaining} 项，详见附件）")
            
            # 插入剩余证书的前5个（图片形式）
            other_certs = [q for q in qualifications if q not in system_certs and q not in credit_certs and q not in honor_certs and q not in partner_certs]
            for cert in other_certs[:5]:
                if cert.get('cert_file'):
                    cert_path = data_dir / cert['cert_file']
                    if cert_path.exists():
                        try:
                            import tempfile
                            with tempfile.TemporaryDirectory() as temp_dir:
                                images = convert_from_path(
                                    str(cert_path),
                                    output_folder=temp_dir,
                                    first_page_only=True,
                                    dpi=200,
                                    fmt='jpg',
                                    use_cropbox=True
                                )
                                
                                if images:
                                    img = PILImage.open(images[0])
                                    img_width, img_height = img.size
                                    
                                    target_width = 500
                                    if img_width > target_width:
                                        ratio = target_width / img_width
                                        new_height = int(img_height * ratio)
                                        img = img.resize((target_width, new_height), PILImage.LANCZOS)
                                        img.save(images[0], quality=85)
                                    
                                    p = doc.add_paragraph()
                                    run = p.add_run(f"• {cert['name']}（{cert['level']}）")
                                    run.bold = True
                                    doc.add_paragraph()
                                    doc.add_picture(images[0], width=Inches(5.5))
                                    doc.add_paragraph()
                        except:
                            pass

        doc.add_page_break()
    """添加类似业绩"""
    p = doc.add_paragraph()
    run = p.add_run("第4章 类似业绩")
    run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        if not cases:
            doc.add_paragraph("（具体业绩清单详见附件）")
            doc.add_page_break()
            return

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ["序号", "项目名称", "客户", "行业", "产品类型", "金额(万元)"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # 填充数据
        for i, case in enumerate(cases[:10], 1):  # 最多显示10个
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = case.get("project_name", "")
            row_cells[2].text = case.get("client", "")
            row_cells[3].text = case.get("industry", "")
            row_cells[4].text = case.get("product_type", "")
            row_cells[5].text = f"{case.get('amount', 0) / 10000:.2f}"

        doc.add_paragraph()
        doc.add_paragraph(f"（其他业绩共 {len(cases)-10} 项，详见附件）")

        doc.add_page_break()

    def _add_after_sales(self, doc: Document, company_info: Dict):
        """添加售后服务"""
        p = doc.add_paragraph()
        run = p.add_run("第5章 售后服务承诺")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        doc.add_paragraph("1. 质量保证")
        doc.add_paragraph("   • 产品符合国家相关标准和规范")
        doc.add_paragraph("   • 提供完整的产品合格证、检验报告等技术资料")
        doc.add_paragraph("   • 质保期：免费保修一年")
        doc.add_paragraph()

        doc.add_paragraph("2. 售后服务")
        doc.add_paragraph("   • 提供7×24小时技术服务热线：400-882-9910")
        doc.add_paragraph(f"   • 技术支持电话：{company_info['phone']}")
        doc.add_paragraph("   • 接到报修后，2小时内响应，24小时内到达现场")
        doc.add_paragraph("   • 提供免费的安装调试、技术培训服务")
        doc.add_paragraph()

        doc.add_paragraph("3. 备品备件")
        doc.add_paragraph("   • 提供质保期内的免费备品备件")
        doc.add_paragraph("   • 质保期后，长期以优惠价格提供备品备件")
        doc.add_paragraph("   • 设备主要元器件采用国内外知名品牌，确保可替换性")
        doc.add_paragraph()

        doc.add_paragraph("4. 技术培训")
        doc.add_paragraph("   • 免费为使用方提供设备操作、维护培训")
        doc.add_paragraph("   • 培训内容包括：设备结构、工作原理、操作规程、故障排除等")
        doc.add_paragraph("   • 培训地点：现场或我司工厂")
        doc.add_paragraph()

        doc.add_page_break()

    def _add_response_commitment(self, doc: Document):
        """
        添加应答承诺函 - 新增
        """
        p = doc.add_paragraph()
        run = p.add_run("第6章 应答承诺函")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        doc.add_paragraph("致：招标人")
        doc.add_paragraph()

        doc.add_paragraph("我公司郑重承诺：")
        doc.add_paragraph()

        commitments = [
            "1. 我公司提供的投标文件完全符合招标文件的各项要求。",
            "2. 我公司具备履行本项目合同的能力，提供的产品符合招标文件的技术要求和质量标准。",
            "3. 我公司保证所提供的产品均为正品，符合国家相关标准，不存在质量问题。",
            "4. 我公司保证严格按照合同约定的交货期交付产品，不影响贵方的工程进度。",
            "5. 我公司保证提供完善的售后服务，确保设备正常运行。",
            "6. 我公司保证报价合理，不存在任何不正当竞争行为。",
            "7. 我公司保证在投标有效期内，不撤销投标文件。",
            "8. 我公司保证遵守招标文件规定的所有条款和要求。"
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_paragraph()

        # 落款
        today = datetime.now()
        doc.add_paragraph(f"承诺单位（盖章）：海越（湖北）电气股份有限公司")
        doc.add_paragraph(f"日期：{today.year}年{today.month}月{today.day}日")

        doc.add_page_break()

    def _add_commitment(self, doc: Document):
        """添加承诺书"""
        p = doc.add_paragraph()
        run = p.add_run("第7章 承诺书")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        doc.add_paragraph("本公司郑重承诺：")
        doc.add_paragraph()

        commitments = [
            "1. 我司具备履行本项目合同的能力，提供的产品符合招标文件的技术要求和质量标准。",
            "2. 我司保证所提供的产品均为正品，符合国家相关标准，不存在质量问题。",
            "3. 我司保证严格按照合同约定的交货期交付产品，不影响贵方的工程进度。",
            "4. 我司保证提供完善的售后服务，确保设备正常运行。",
            "5. 我司保证报价合理，不存在任何不正当竞争行为。",
            "6. 我司保证在投标有效期内，不撤销投标文件。",
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_paragraph()

        # 落款
        today = datetime.now()
        doc.add_paragraph(f"承诺单位（盖章）：海越（湖北）电气股份有限公司")
        doc.add_paragraph(f"日期：{today.year}年{today.month}月{today.day}日")

    def _add_tech_commitment(self, doc: Document):
        """添加技术承诺书"""
        p = doc.add_paragraph()
        run = p.add_run("技术承诺书")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        doc.add_paragraph("本公司郑重承诺：")
        doc.add_paragraph()

        commitments = [
            "1. 我司提供的设备技术方案完全符合招标文件的技术要求。",
            "2. 我司采用的设备和元器件均符合国家标准和行业规范。",
            "3. 我司保证设备的技术参数真实有效，不存在虚假陈述。",
            "4. 我司愿意配合贵方进行技术澄清和方案优化。",
            "5. 我司提供的技术资料完整、准确、规范。",
            "6. 我司提供的技术服务专业、及时、高效。",
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_paragraph()

        # 落款
        today = datetime.now()
        doc.add_paragraph(f"承诺单位（盖章）：海越（湖北）电气股份有限公司")
        doc.add_paragraph(f"日期：{today.year}年{today.month}月{today.day}日")

    def _add_commercial_commitment(self, doc: Document):
        """添加商务承诺书"""
        p = doc.add_paragraph()
        run = p.add_run("商务承诺书")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"

        doc.add_paragraph()

        doc.add_paragraph("本公司郑重承诺：")
        doc.add_paragraph()

        commitments = [
            "1. 我司报价合理，不存在任何不正当竞争行为。",
            "2. 我司保证在投标有效期内，不撤销投标文件。",
            "3. 我司严格按照合同约定的交货期交付产品，不影响贵方的工程进度。",
            "4. 我司提供完善的售后服务，确保设备正常运行。",
            "5. 我司愿意接受贵方对价格、交货期、服务条款的合理要求。",
            "6. 我司遵守招标文件规定的所有商务条款。",
        ]

        for commitment in commitments:
            doc.add_paragraph(commitment)

        doc.add_paragraph()

        # 落款
        today = datetime.now()
        doc.add_paragraph(f"承诺单位（盖章）：海越（湖北）电气股份有限公司")
        doc.add_paragraph(f"日期：{today.year}年{today.month}月{today.day}日")

    def _nsdecls(self, prefix):
        """返回XML命名空间声明"""
        return f'xmlns:w="{prefix}"'

    def _parse_xml(self, xml_str):
        """解析XML字符串"""
        return OxmlElement(xml_str)


# 测试代码
if __name__ == "__main__":
    from config import COMPANY_INFO

    generator = BidDocumentGenerator(
        templates_dir=Path(__file__).parent / "templates",
        output_dir=Path(__file__).parent / "output"
    )

    # 模拟数据
    tender_info = {
        "project_info": {
            "project_name": "测试项目10kV开关柜采购",
            "project_no": "TEST-2025-001",
            "tenderer": "测试单位",
        },
        "product_requirements": ["10kV开关柜", "低压开关柜"],
    }

    matched_data = {
        "qualifications": [
            {"name": "质量管理体系认证", "level": "ISO9001", "cert_no": "CN14Q46720ROM", "valid_until": "2025-12-31"},
            {"name": "AAA级信用企业", "level": "AAA级", "cert_no": "", "valid_until": ""},
        ],
        "cases": [
            {"project_name": "测试项目1", "client": "测试客户1", "industry": "钢铁", "product_type": "高压开关柜", "amount": 1000000},
            {"project_name": "测试项目2", "client": "测试客户2", "industry": "环保", "product_type": "低压开关柜", "amount": 500000},
        ],
    }

    quote_data = {
        "products": [
            {"产品名称": "10kV开关柜", "型号/规格": "KYN28A-12", "单位": "台", "数量": 10, "单价": 50000},
            {"产品名称": "低压开关柜", "型号/规格": "MNS", "单位": "台", "数量": 5, "单价": 30000},
        ],
    }

    # 生成投标文件
    output_path = generator.generate_bid(tender_info, COMPANY_INFO, matched_data, quote_data)
    print(f"投标文件已生成: {output_path}")
